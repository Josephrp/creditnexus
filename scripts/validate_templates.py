"""
Template validation tool to audit template files for placeholders.

This script:
1. Scans all template files in storage/templates
2. Extracts all placeholders using regex patterns
3. Validates placeholder format consistency
4. Checks for unmapped placeholders (not in field mappings)
5. Reports issues and recommendations
"""

import sys
import re
from pathlib import Path
from typing import List, Dict, Set, Tuple
from docx import Document

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.db import SessionLocal, init_db
from app.templates.registry import TemplateRegistry
from app.db.models import LMATemplate, TemplateFieldMapping

# Placeholder patterns (matching renderer patterns)
BRACKET_PATTERN = re.compile(r'\[([A-Z_][A-Z0-9_]*)\]')
DOUBLE_CURLY_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
SINGLE_CURLY_PATTERN = re.compile(r'\{([^}]+)\}')

logger = None


def extract_placeholders_from_docx(file_path: Path) -> Dict[str, List[str]]:
    """
    Extract all placeholders from a Word document.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Dictionary with placeholder types as keys and lists of placeholders as values
    """
    placeholders = {
        'bracket': [],      # [FIELD_NAME]
        'double_curly': [], # {{field_name}}
        'single_curly': [],  # {field_name}
        'all': []           # All unique placeholders
    }
    
    try:
        doc = Document(file_path)
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            placeholders['bracket'].extend(BRACKET_PATTERN.findall(text))
            placeholders['double_curly'].extend(DOUBLE_CURLY_PATTERN.findall(text))
            placeholders['single_curly'].extend(SINGLE_CURLY_PATTERN.findall(text))
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        placeholders['bracket'].extend(BRACKET_PATTERN.findall(text))
                        placeholders['double_curly'].extend(DOUBLE_CURLY_PATTERN.findall(text))
                        placeholders['single_curly'].extend(SINGLE_CURLY_PATTERN.findall(text))
        
        # Extract from headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        text = paragraph.text
                        placeholders['bracket'].extend(BRACKET_PATTERN.findall(text))
                        placeholders['double_curly'].extend(DOUBLE_CURLY_PATTERN.findall(text))
                        placeholders['single_curly'].extend(SINGLE_CURLY_PATTERN.findall(text))
            
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        text = paragraph.text
                        placeholders['bracket'].extend(BRACKET_PATTERN.findall(text))
                        placeholders['double_curly'].extend(DOUBLE_CURLY_PATTERN.findall(text))
                        placeholders['single_curly'].extend(SINGLE_CURLY_PATTERN.findall(text))
        
        # Normalize and deduplicate
        # Convert bracket placeholders to [FIELD_NAME] format
        bracket_normalized = [f"[{p}]" for p in set(placeholders['bracket'])]
        # Convert double curly to bracket format for comparison
        double_curly_normalized = [f"[{p.upper().replace('.', '_').replace('-', '_')}]" for p in set(placeholders['double_curly'])]
        # Single curly are typically CDM paths, keep as-is
        single_curly_unique = list(set(placeholders['single_curly']))
        
        placeholders['bracket'] = bracket_normalized
        placeholders['double_curly'] = double_curly_normalized
        placeholders['single_curly'] = single_curly_unique
        placeholders['all'] = list(set(bracket_normalized + double_curly_normalized + single_curly_unique))
        
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return placeholders
    
    return placeholders


def get_template_mappings(db, template_id: int) -> Tuple[Set[str], Set[str]]:
    """
    Get all mapped template field names for a template.
    
    Args:
        db: Database session
        template_id: Template ID
        
    Returns:
        Tuple of (mapped_fields, ai_generated_fields) sets
    """
    mappings = TemplateRegistry.get_field_mappings(db, template_id)
    mapped_fields = {m.template_field for m in mappings if m.mapping_type != "ai_generated"}
    ai_generated_fields = {m.template_field for m in mappings if m.mapping_type == "ai_generated"}
    
    # Also check template metadata for AI-generated sections
    template = db.query(LMATemplate).filter(LMATemplate.id == template_id).first()
    if template and template.ai_generated_sections:
        # Convert section names to placeholder format
        for section in template.ai_generated_sections:
            placeholder = f"[{section.upper().replace(' ', '_')}]"
            ai_generated_fields.add(placeholder)
    
    return mapped_fields, ai_generated_fields


def validate_template(template: LMATemplate, db) -> Dict[str, any]:
    """
    Validate a single template file.
    
    Args:
        template: LMATemplate instance
        db: Database session
        
    Returns:
        Validation results dictionary
    """
    template_path = Path(template.file_path)
    
    if not template_path.exists():
        return {
            'template_code': template.template_code,
            'status': 'error',
            'error': f'Template file not found: {template_path}',
            'placeholders': {},
            'issues': []
        }
    
    # Extract placeholders
    placeholders = extract_placeholders_from_docx(template_path)
    
    # Get field mappings (separate direct mappings from AI-generated)
    mapped_fields, ai_generated_fields = get_template_mappings(db, template.id)
    
    # Find unmapped placeholders (exclude AI-generated fields)
    bracket_placeholders = set(placeholders['bracket'])
    unmapped = bracket_placeholders - mapped_fields - ai_generated_fields
    
    # Identify issues
    issues = []
    
    # Check for double curly placeholders (should be converted to bracket format)
    if placeholders['double_curly']:
        issues.append({
            'type': 'format_inconsistency',
            'severity': 'warning',
            'message': f"Found {len(placeholders['double_curly'])} double curly placeholders ({{{{field}}}}). "
                      "These should be converted to bracket format [FIELD_NAME] for consistency.",
            'placeholders': placeholders['double_curly'][:10]  # First 10
        })
    
    # Check for unmapped bracket placeholders
    if unmapped:
        issues.append({
            'type': 'unmapped_placeholder',
            'severity': 'error',
            'message': f"Found {len(unmapped)} unmapped placeholders. These need field mappings.",
            'placeholders': list(unmapped)
        })
    
    # Check for single curly placeholders (CDM paths - these are OK)
    if placeholders['single_curly']:
        issues.append({
            'type': 'cdm_path_placeholder',
            'severity': 'info',
            'message': f"Found {len(placeholders['single_curly'])} CDM path placeholders ({{field}}). "
                      "These are evaluated directly from CDM data.",
            'placeholders': placeholders['single_curly'][:10]  # First 10
        })
    
    # Check if template has no placeholders (might be an issue)
    if not placeholders['all']:
        issues.append({
            'type': 'no_placeholders',
            'severity': 'warning',
            'message': 'Template has no placeholders. This might be a static template.',
            'placeholders': []
        })
    
    return {
        'template_code': template.template_code,
        'template_name': template.name,
        'status': 'error' if any(i['severity'] == 'error' for i in issues) else 'warning' if issues else 'ok',
        'placeholders': {
            'bracket_count': len(placeholders['bracket']),
            'double_curly_count': len(placeholders['double_curly']),
            'single_curly_count': len(placeholders['single_curly']),
            'total_unique': len(placeholders['all']),
            'mapped_count': len(bracket_placeholders & mapped_fields),
            'ai_generated_count': len(bracket_placeholders & ai_generated_fields),
            'unmapped_count': len(unmapped)
        },
        'issues': issues,
        'file_path': str(template_path)
    }


def main():
    """Validate all templates in the database."""
    # Initialize database
    init_db()
    db = SessionLocal()
    
    try:
        # Get all templates
        templates = db.query(LMATemplate).all()
        
        if not templates:
            print("No templates found in database.")
            return
        
        print(f"Validating {len(templates)} template(s)...\n")
        print("=" * 80)
        
        results = []
        total_issues = {'error': 0, 'warning': 0, 'info': 0}
        
        for template in templates:
            result = validate_template(template, db)
            results.append(result)
            
            # Print template summary
            print(f"\nTemplate: {result['template_code']} - {result['template_name']}")
            print(f"Status: {result['status'].upper()}")
            print(f"File: {result.get('file_path', 'N/A')}")
            
            if 'error' in result:
                print(f"ERROR: {result['error']}")
                continue
            
            # Print placeholder summary
            ph = result['placeholders']
            print(f"\nPlaceholders:")
            print(f"  Bracket format [FIELD]: {ph['bracket_count']}")
            print(f"  Double curly {{field}}: {ph['double_curly_count']}")
            print(f"  Single curly {{field}}: {ph['single_curly_count']}")
            print(f"  Total unique: {ph['total_unique']}")
            print(f"  Mapped (direct): {ph['mapped_count']}")
            print(f"  AI-generated: {ph.get('ai_generated_count', 0)}")
            print(f"  Unmapped: {ph['unmapped_count']}")
            
            # Print issues
            if result['issues']:
                print(f"\nIssues ({len(result['issues'])}):")
                for issue in result['issues']:
                    severity_icon = {
                        'error': '[ERROR]',
                        'warning': '[WARN]',
                        'info': '[INFO]'
                    }.get(issue['severity'], '[?]')
                    print(f"  {severity_icon} {issue['message']}")
                    if issue.get('placeholders'):
                        print(f"    Examples: {', '.join(issue['placeholders'][:5])}")
                    total_issues[issue['severity']] += 1
            else:
                print("\n[OK] No issues found!")
            
            print("-" * 80)
        
        # Print summary
        print(f"\n{'=' * 80}")
        print("VALIDATION SUMMARY")
        print(f"{'=' * 80}")
        print(f"Templates validated: {len(templates)}")
        print(f"Templates with errors: {sum(1 for r in results if r['status'] == 'error')}")
        print(f"Templates with warnings: {sum(1 for r in results if 'warning' in str(r.get('issues', [])))}")
        print(f"Templates OK: {sum(1 for r in results if r['status'] == 'ok')}")
        print(f"\nTotal issues:")
        print(f"  Errors: {total_issues['error']}")
        print(f"  Warnings: {total_issues['warning']}")
        print(f"  Info: {total_issues['info']}")
        
        # Recommendations
        if total_issues['error'] > 0 or total_issues['warning'] > 0:
            print(f"\n{'=' * 80}")
            print("RECOMMENDATIONS")
            print(f"{'=' * 80}")
            
            if total_issues['error'] > 0:
                print("1. Create field mappings for all unmapped placeholders")
                print("2. Run: python scripts/seed_field_mappings.py")
            
            if any('double_curly' in str(r.get('issues', [])) for r in results):
                print("3. Convert double curly placeholders {{field}} to bracket format [FIELD_NAME]")
                print("4. Update template files or regeneration script")
            
            print("5. Ensure all required fields are extracted by the extraction prompt")
        
    except Exception as e:
        import traceback
        print(f"Error during validation: {e}")
        traceback.print_exc()
        sys.exit(1)
    finally:
        db.close()


if __name__ == "__main__":
    main()
