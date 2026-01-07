"""
Script to generate dummy LMA template files with placeholders.

Creates Word documents with placeholder tags for testing template generation.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))


def generate_dummy_facility_agreement(output_path: Path):
    """
    Generate dummy Facility Agreement template.
    
    Args:
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('FACILITY AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parties Section
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('LEI: [BORROWER_LEI]')
    doc.add_paragraph('Lenders: [LENDERS_LIST]')
    doc.add_paragraph('Administrative Agent: [ADMIN_AGENT_NAME]')
    
    # Facility Details
    doc.add_heading('FACILITY DETAILS', 1)
    doc.add_paragraph(f'Facility Name: [FACILITY_NAME]')
    doc.add_paragraph(f'Total Commitment: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph(f'Maturity Date: [MATURITY_DATE]')
    doc.add_paragraph(f'Agreement Date: [AGREEMENT_DATE]')
    
    # Interest Terms
    doc.add_heading('INTEREST TERMS', 1)
    doc.add_paragraph(f'Benchmark: [BENCHMARK]')
    doc.add_paragraph(f'Spread: [SPREAD] basis points')
    doc.add_paragraph(f'Payment Frequency: [PAYMENT_FREQUENCY]')
    
    # Governing Law
    doc.add_heading('GOVERNING LAW', 1)
    doc.add_paragraph(f'This Agreement is governed by [GOVERNING_LAW] law.')
    
    # AI-Generated Sections
    doc.add_heading('REPRESENTATIONS AND WARRANTIES', 1)
    doc.add_paragraph('[REPRESENTATIONS_AND_WARRANTIES]')
    
    doc.add_heading('CONDITIONS PRECEDENT', 1)
    doc.add_paragraph('[CONDITIONS_PRECEDENT]')
    
    doc.add_heading('COVENANTS', 1)
    doc.add_paragraph('[COVENANTS]')
    
    doc.add_heading('EVENTS OF DEFAULT', 1)
    doc.add_paragraph('[EVENTS_OF_DEFAULT]')
    
    # ESG Section (if applicable)
    doc.add_heading('SUSTAINABILITY-LINKED LOAN PROVISIONS', 1)
    doc.add_paragraph('Sustainability-Linked: [SUSTAINABILITY_LINKED]')
    doc.add_paragraph('ESG KPI Targets: [ESG_KPI_TARGETS]')
    doc.add_paragraph('SPT Definitions: [SPT_DEFINITIONS]')
    doc.add_paragraph('Margin Adjustment: [MARGIN_ADJUSTMENT]')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated Facility Agreement template: {output_path}")


def generate_dummy_term_sheet(output_path: Path):
    """
    Generate dummy Term Sheet template.
    
    Args:
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('TERM SHEET', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Borrower
    doc.add_heading('BORROWER', 1)
    doc.add_paragraph('[BORROWER]')
    
    # Facility Type
    doc.add_heading('FACILITY TYPE', 1)
    doc.add_paragraph('[FACILITY_TYPE]')
    
    # Total Commitment
    doc.add_heading('TOTAL COMMITMENT', 1)
    doc.add_paragraph('[TOTAL_COMMITMENT] [CURRENCY]')
    
    # Pricing
    doc.add_heading('PRICING', 1)
    doc.add_paragraph('[PRICING]')
    
    # Maturity
    doc.add_heading('MATURITY', 1)
    doc.add_paragraph('[MATURITY_DATE]')
    
    # Purpose (AI-generated)
    doc.add_heading('PURPOSE', 1)
    doc.add_paragraph('[PURPOSE]')
    
    # Conditions Precedent (AI-generated)
    doc.add_heading('CONDITIONS PRECEDENT', 1)
    doc.add_paragraph('[CONDITIONS_PRECEDENT]')
    
    # Representations (AI-generated)
    doc.add_heading('REPRESENTATIONS', 1)
    doc.add_paragraph('[REPRESENTATIONS]')
    
    # Fees (AI-generated)
    doc.add_heading('FEES', 1)
    doc.add_paragraph('[FEES]')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated Term Sheet template: {output_path}")


def generate_dummy_confidentiality_agreement(output_path: Path):
    """
    Generate dummy Confidentiality Agreement template.
    
    Args:
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('CONFIDENTIALITY AND NO FRONT RUNNING AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parties
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Deal ID: [DEAL_ID]')
    
    # Confidentiality Obligations (AI-generated)
    doc.add_heading('CONFIDENTIALITY OBLIGATIONS', 1)
    doc.add_paragraph('[CONFIDENTIALITY_OBLIGATIONS]')
    
    # No Front Running Undertaking (AI-generated)
    doc.add_heading('NO FRONT RUNNING UNDERTAKING', 1)
    doc.add_paragraph('[NO_FRONT_RUNNING_UNDERTAKING]')
    
    # Permitted Disclosures (AI-generated)
    doc.add_heading('PERMITTED DISCLOSURES', 1)
    doc.add_paragraph('[PERMITTED_DISCLOSURES]')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated Confidentiality Agreement template: {output_path}")


def generate_dummy_ref_facility_agreement(output_path: Path):
    """
    Generate dummy Real Estate Finance (REF) Facility Agreement template.
    
    Args:
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('REAL ESTATE FINANCE FACILITY AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parties
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Lenders: [LENDERS_LIST]')
    
    # Facility Details
    doc.add_heading('FACILITY DETAILS', 1)
    doc.add_paragraph(f'Facility Name: [FACILITY_NAME]')
    doc.add_paragraph(f'Total Commitment: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph(f'Maturity Date: [MATURITY_DATE]')
    
    # Property Description (AI-generated)
    doc.add_heading('PROPERTY DESCRIPTION', 1)
    doc.add_paragraph('[PROPERTY_DESCRIPTION]')
    
    # Security Package (AI-generated)
    doc.add_heading('SECURITY PACKAGE', 1)
    doc.add_paragraph('[SECURITY_PACKAGE]')
    
    # Valuation Requirements (AI-generated)
    doc.add_heading('VALUATION REQUIREMENTS', 1)
    doc.add_paragraph('[VALUATION_REQUIREMENTS]')
    
    # Representations and Warranties (AI-generated)
    doc.add_heading('REPRESENTATIONS AND WARRANTIES', 1)
    doc.add_paragraph('[REPRESENTATIONS_AND_WARRANTIES]')
    
    # Conditions Precedent (AI-generated)
    doc.add_heading('CONDITIONS PRECEDENT', 1)
    doc.add_paragraph('[CONDITIONS_PRECEDENT]')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated REF Facility Agreement template: {output_path}")


def generate_dummy_sll_facility_agreement(output_path: Path):
    """
    Generate dummy Sustainability-Linked Loan (SLL) Facility Agreement template.
    
    Args:
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('SUSTAINABILITY-LINKED LOAN FACILITY AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Parties
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Lenders: [LENDERS_LIST]')
    
    # Facility Details
    doc.add_heading('FACILITY DETAILS', 1)
    doc.add_paragraph(f'Facility Name: [FACILITY_NAME]')
    doc.add_paragraph(f'Total Commitment: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph(f'Maturity Date: [MATURITY_DATE]')
    doc.add_paragraph(f'Sustainability-Linked: [SUSTAINABILITY_LINKED]')
    
    # SPT Definitions (AI-generated)
    doc.add_heading('SUSTAINABILITY PERFORMANCE TARGETS (SPT)', 1)
    doc.add_paragraph('[SPT_DEFINITIONS]')
    
    # SPT Measurement Methodology (AI-generated)
    doc.add_heading('SPT MEASUREMENT METHODOLOGY', 1)
    doc.add_paragraph('[SPT_MEASUREMENT_METHODOLOGY]')
    
    # Margin Adjustment Mechanism (AI-generated)
    doc.add_heading('MARGIN ADJUSTMENT MECHANISM', 1)
    doc.add_paragraph('[MARGIN_ADJUSTMENT_MECHANISM]')
    
    # Reporting Requirements (AI-generated)
    doc.add_heading('REPORTING REQUIREMENTS', 1)
    doc.add_paragraph('[REPORTING_REQUIREMENTS]')
    
    # Verification Process (AI-generated)
    doc.add_heading('VERIFICATION PROCESS', 1)
    doc.add_paragraph('[VERIFICATION_PROCESS]')
    
    # ESG KPI Targets
    doc.add_heading('ESG KPI TARGETS', 1)
    doc.add_paragraph('[ESG_KPI_TARGETS]')
    
    # Standard Sections
    doc.add_heading('REPRESENTATIONS AND WARRANTIES', 1)
    doc.add_paragraph('[REPRESENTATIONS_AND_WARRANTIES]')
    
    doc.add_heading('COVENANTS', 1)
    doc.add_paragraph('[COVENANTS]')
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Generated SLL Facility Agreement template: {output_path}")


def main():
    """Main execution: generate all dummy templates."""
    # Create directory structure
    base_path = Path("storage/templates")
    base_path.mkdir(parents=True, exist_ok=True)
    
    # Generate templates
    templates = [
        ("facility_agreements/corporate_lending/english/CL-FA-EN-2024.1.docx", generate_dummy_facility_agreement),
        ("term_sheets/corporate_lending/english/CL-TS-EN-2024.1.docx", generate_dummy_term_sheet),
        ("confidentiality/primary_syndication/CONF-PS-2024.1.docx", generate_dummy_confidentiality_agreement),
        ("facility_agreements/real_estate/english/REF-FA-EN-2024.1.docx", generate_dummy_ref_facility_agreement),
        ("facility_agreements/sustainable/sll/english/SLL-FA-EN-2024.1.docx", generate_dummy_sll_facility_agreement),
    ]
    
    for rel_path, generator_func in templates:
        output_path = base_path / rel_path
        generator_func(output_path)
    
    print(f"\nSuccessfully generated {len(templates)} dummy template(s) in {base_path}")


if __name__ == "__main__":
    main()












