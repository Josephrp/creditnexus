"""
Script to create template Word files for LMA templates.

Creates placeholder Word documents from templates_metadata.json or database.
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List, Optional

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.db import SessionLocal, init_db
from app.templates.registry import TemplateRegistry


def _normalize_field_name(field_path: str) -> str:
    """
    Convert CDM field path to placeholder format.
    
    Examples:
    - "parties[role='Borrower'].name" -> "BORROWER_NAME"
    - "facilities[0].facility_name" -> "FACILITY_NAME"
    - "facilities[0].commitment_amount.amount" -> "COMMITMENT_AMOUNT"
    """
    # Extract role from filter if present
    if "role=" in field_path:
        role_match = field_path.split("role=")[1].split("'")[1] if "'" in field_path.split("role=")[1] else None
        if role_match:
            field_name = field_path.split(".")[-1] if "." in field_path else field_path
            return f"{role_match.upper()}_{field_name.upper()}"
    
    # Extract last part of path
    last_part = field_path.split(".")[-1] if "." in field_path else field_path
    # Remove array indices
    last_part = last_part.split("[")[0] if "[" in last_part else last_part
    # Convert to uppercase with underscores
    return last_part.upper().replace("-", "_")


def _normalize_ai_section(section: str) -> str:
    """Convert AI section name to placeholder format."""
    return section.upper().replace("-", "_")


def create_template_file_from_metadata(template_data: Dict[str, Any], output_path: Path):
    """
    Create a placeholder Word template file from metadata.
    
    Args:
        template_data: Template metadata dictionary
        output_path: Path to save the template file
    """
    doc = Document()
    
    # Title based on template name
    title_text = template_data.get("name", "DOCUMENT")
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Template Info
    doc.add_paragraph(f'Template Code: {template_data.get("template_code", "N/A")}')
    doc.add_paragraph(f'Version: {template_data.get("version", "N/A")}')
    doc.add_paragraph(f'Category: {template_data.get("category", "N/A")}')
    if template_data.get("governing_law"):
        doc.add_paragraph(f'Governing Law: {template_data.get("governing_law")}')
    doc.add_paragraph('')
    
    # Determine template structure based on category
    category = template_data.get("category", "").lower()
    
    # Common sections for most templates
    if "facility agreement" in category or "term sheet" in category:
        _add_facility_agreement_sections(doc, template_data)
    elif "sustainable finance" in category or "sustainability" in category.lower():
        _add_sustainability_loan_sections(doc, template_data)
    elif "confidentiality" in category.lower():
        _add_confidentiality_agreement_sections(doc, template_data)
    elif "secondary trading" in category.lower():
        _add_secondary_trading_sections(doc, template_data)
    elif "intercreditor" in category.lower() or "security" in category.lower():
        _add_intercreditor_sections(doc, template_data)
    elif "regulatory" in category.lower():
        _add_regulatory_sections(doc, template_data)
    elif "restructuring" in category.lower():
        _add_restructuring_sections(doc, template_data)
    elif "supporting" in category.lower() or "opinion" in template_data.get("name", "").lower():
        _add_legal_opinion_sections(doc, template_data)
    elif "regional" in category.lower() or "consumer" in category.lower():
        _add_consumer_credit_sections(doc, template_data)
    elif "origination" in category.lower():
        _add_origination_sections(doc, template_data)
    else:
        # Generic template structure
        _add_generic_sections(doc, template_data)
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created template file: {output_path}")


def _add_facility_agreement_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for facility agreement templates."""
    # Parties Section
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    if "parties[role='Borrower'].lei" in template_data.get("required_cdm_fields", []):
        doc.add_paragraph('Borrower LEI: [BORROWER_LEI]')
    doc.add_paragraph('Lenders: [LENDER_NAME]')
    doc.add_paragraph('Administrative Agent: [ADMINISTRATIVE_AGENT_NAME]')
    doc.add_paragraph('')
    
    # Facility Details
    doc.add_heading('FACILITY DETAILS', 1)
    doc.add_paragraph('Facility Name: [FACILITY_NAME]')
    doc.add_paragraph('Total Commitment: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph('Maturity Date: [MATURITY_DATE]')
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # Interest Terms
    doc.add_heading('INTEREST TERMS', 1)
    doc.add_paragraph('Benchmark: [BENCHMARK]')
    doc.add_paragraph('Spread: [SPREAD] basis points')
    doc.add_paragraph('Payment Frequency: [PAYMENT_FREQUENCY]')
    doc.add_paragraph('')
    
    # Governing Law
    doc.add_heading('GOVERNING LAW', 1)
    doc.add_paragraph('This Agreement is governed by [GOVERNING_LAW] law.')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_sustainability_loan_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for sustainability-linked loan templates."""
    _add_facility_agreement_sections(doc, template_data)
    
    # ESG Section
    doc.add_heading('SUSTAINABILITY-LINKED LOAN PROVISIONS', 1)
    doc.add_paragraph('Sustainability-Linked: [SUSTAINABILITY_LINKED]')
    doc.add_paragraph('ESG KPI Targets: [ESG_KPI_TARGETS]')
    if "spt_definitions" in str(template_data.get("optional_cdm_fields", [])):
        doc.add_paragraph('SPT Definitions: [SPT_DEFINITIONS]')
    if "margin_adjustment" in str(template_data.get("optional_cdm_fields", [])):
        doc.add_paragraph('Margin Adjustment: [MARGIN_ADJUSTMENT]')
    doc.add_paragraph('')
    
    # AI-Generated ESG Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        if "sustainability" in section.lower() or "kpi" in section.lower() or "esg" in section.lower():
            section_title = section.replace("_", " ").title()
            doc.add_heading(section_title.upper(), 1)
            placeholder = f"[{_normalize_ai_section(section)}]"
            doc.add_paragraph(placeholder)
            doc.add_paragraph('')


def _add_confidentiality_agreement_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for confidentiality agreement templates."""
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Disclosing Party: [BORROWER_NAME]')
    doc.add_paragraph('Receiving Party: [LENDER_NAME]')
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_secondary_trading_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for secondary trading templates."""
    doc.add_heading('TRADE DETAILS', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Facility: [FACILITY_NAME]')
    doc.add_paragraph('Trade Amount: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph('Trade Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_intercreditor_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for intercreditor agreement templates."""
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Senior Lenders: [LENDER_NAME]')
    doc.add_paragraph('Administrative Agent: [ADMINISTRATIVE_AGENT_NAME]')
    doc.add_paragraph('Facility: [FACILITY_NAME]')
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_regulatory_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for regulatory compliance templates."""
    doc.add_heading('PARTIES', 1)
    doc.add_paragraph('Entity Name: [BORROWER_NAME]')
    if "parties[role='Borrower'].lei" in template_data.get("required_cdm_fields", []):
        doc.add_paragraph('LEI: [BORROWER_LEI]')
    doc.add_paragraph('Certification Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_restructuring_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for restructuring templates."""
    doc.add_heading('RESTRUCTURING DETAILS', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Facility: [FACILITY_NAME]')
    doc.add_paragraph('Outstanding Amount: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_legal_opinion_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for legal opinion templates."""
    doc.add_heading('OPINION DETAILS', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Facility: [FACILITY_NAME]')
    doc.add_paragraph('Governing Law: [GOVERNING_LAW]')
    doc.add_paragraph('Opinion Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_consumer_credit_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for consumer credit templates."""
    doc.add_heading('CREDIT AGREEMENT DETAILS', 1)
    doc.add_paragraph('Borrower: [BORROWER_NAME]')
    doc.add_paragraph('Lender: [LENDER_NAME]')
    doc.add_paragraph('Credit Amount: [COMMITMENT_AMOUNT] [CURRENCY]')
    doc.add_paragraph('Term: [MATURITY_DATE]')
    doc.add_paragraph('Interest Rate: [BENCHMARK] + [SPREAD] basis points')
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_origination_sections(doc: Document, template_data: Dict[str, Any]):
    """Add sections for origination document templates."""
    doc.add_heading('CERTIFICATION DETAILS', 1)
    doc.add_paragraph('Entity Name: [BORROWER_NAME]')
    if "parties[role='Borrower'].lei" in template_data.get("required_cdm_fields", []):
        doc.add_paragraph('LEI: [BORROWER_LEI]')
    doc.add_paragraph('Certification Date: [AGREEMENT_DATE]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def _add_generic_sections(doc: Document, template_data: Dict[str, Any]):
    """Add generic sections for unknown template types."""
    doc.add_heading('DOCUMENT DETAILS', 1)
    doc.add_paragraph('Agreement Date: [AGREEMENT_DATE]')
    if "parties[role='Borrower'].name" in template_data.get("required_cdm_fields", []):
        doc.add_paragraph('Borrower: [BORROWER_NAME]')
    if "parties[role='Lender'].name" in template_data.get("required_cdm_fields", []):
        doc.add_paragraph('Lender: [LENDER_NAME]')
    doc.add_paragraph('')
    
    # AI-Generated Sections
    ai_sections = template_data.get("ai_generated_sections", [])
    for section in ai_sections:
        section_title = section.replace("_", " ").title()
        doc.add_heading(section_title.upper(), 1)
        placeholder = f"[{_normalize_ai_section(section)}]"
        doc.add_paragraph(placeholder)
        doc.add_paragraph('')


def load_template_metadata(json_path: Path) -> List[Dict[str, Any]]:
    """
    Load template metadata from JSON file.
    
    Args:
        json_path: Path to templates_metadata.json
        
    Returns:
        List of template metadata dictionaries
    """
    if not json_path.exists():
        return []
    
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
            elif isinstance(data, dict) and "templates" in data:
                return data["templates"]
            else:
                return []
    except Exception as e:
        print(f"Error loading template metadata from {json_path}: {e}")
        return []


def main(use_metadata: bool = True, force_regenerate: bool = False):
    """
    Create template files from metadata or database.
    
    Args:
        use_metadata: If True, load from templates_metadata.json; if False, use database
        force_regenerate: If True, overwrite existing template files
    """
    # Initialize database
    init_db()
    
    templates_data = []
    
    if use_metadata:
        # Load from metadata file
        json_paths = [
            Path("data/templates_metadata.json"),
            Path("scripts/templates_metadata.json"),
            Path("storage/templates_metadata.json"),
        ]
        
        for json_path in json_paths:
            if json_path.exists():
                templates_data = load_template_metadata(json_path)
                print(f"Loaded {len(templates_data)} template(s) from {json_path}")
                break
        
        if not templates_data:
            print("No template metadata file found. Falling back to database...")
            use_metadata = False
    
    if not use_metadata:
        # Load from database
        db = SessionLocal()
        try:
            templates = TemplateRegistry.list_templates(db)
            if not templates:
                print("No templates found in database. Please seed templates first.")
                return
            
            # Convert database templates to metadata format
            for template in templates:
                templates_data.append({
                    "template_code": template.template_code,
                    "name": template.name,
                    "category": template.category,
                    "subcategory": template.subcategory,
                    "governing_law": template.governing_law,
                    "version": template.version,
                    "file_path": template.file_path,
                    "required_cdm_fields": template.required_fields or [],
                    "optional_cdm_fields": template.optional_fields or [],
                    "ai_generated_sections": template.ai_generated_sections or [],
                })
            print(f"Loaded {len(templates_data)} template(s) from database")
        finally:
            db.close()
    
    if not templates_data:
        print("No templates found. Exiting.")
        return
    
    # Base path for templates
    base_path = Path("storage/templates")
    base_path.mkdir(parents=True, exist_ok=True)
    
    created_count = 0
    
    for template_data in templates_data:
        template_code = template_data.get("template_code")
        version = template_data.get("version", "2024.1")
        
        # Determine output path
        if template_data.get("file_path"):
            output_path = Path(template_data["file_path"])
        else:
            filename = f"{template_code}-{version}.docx"
            output_path = base_path / filename
        
        # Force regeneration - remove existing file if it exists
        if force_regenerate and output_path.exists():
            output_path.unlink()
            print(f"Removed existing template file: {output_path}")
        
        # Skip if file exists and not forcing regeneration
        if output_path.exists() and not force_regenerate:
            print(f"Skipping existing template file: {output_path}")
            continue
        
        # Create template file
        try:
            create_template_file_from_metadata(template_data, output_path)
            created_count += 1
        except Exception as e:
            print(f"Error creating template file for {template_code}: {e}")
            continue
    
    print(f"\nSuccessfully created {created_count} template file(s)")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Create LMA template Word files")
    parser.add_argument(
        "--use-db",
        action="store_true",
        help="Load templates from database instead of metadata file"
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Force regeneration of existing template files"
    )
    
    args = parser.parse_args()
    main(use_metadata=not args.use_db, force_regenerate=args.force)
