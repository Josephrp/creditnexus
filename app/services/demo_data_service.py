"""
Demo Data Seeding Service for CreditNexus.

This service provides a unified interface for seeding demo data including:
- Users (via existing seed scripts)
- Templates (via existing seed scripts)
- Policies (via existing seed scripts)
- AI-generated deals with complete data flow
- Documents, workflows, and related entities
"""

import logging
import random
import json
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Callable
from datetime import datetime, timedelta, date
from decimal import Decimal
from pathlib import Path
from sqlalchemy.orm import Session
from sqlalchemy import and_

from app.db.models import (
    User, UserRole, Deal, Application, Document, DocumentVersion, Workflow,
    GeneratedDocument, DealNote, PolicyDecision, DealStatus,
    DealType, ApplicationStatus, ApplicationType, WorkflowState, GeneratedDocumentStatus,
    GreenFinanceAssessment, SecuritizationPool, SecuritizationTranche,
    SecuritizationPoolAsset, RegulatoryFiling, NotarizationRecord
)
from app.models.cdm import CreditAgreement
from app.models.loan_asset import LoanAsset
from app.models.user_profile import UserProfileData
from app.core.config import settings
from app.services.file_storage_service import FileStorageService

# #region agent log
log_data = {
    "sessionId": "debug-session",
    "runId": "post-fix",
    "hypothesisId": "A",
    "location": "demo_data_service.py:33",
    "message": "Post-fix: Settings import verification",
    "data": {
        "settings_in_globals": "settings" in globals(),
        "settings_type": type(settings).__name__ if "settings" in globals() else "NOT_FOUND",
        "has_enhanced_satellite_attr": hasattr(settings, "ENHANCED_SATELLITE_ENABLED") if "settings" in globals() else False,
        "enhanced_satellite_value": getattr(settings, "ENHANCED_SATELLITE_ENABLED", None) if "settings" in globals() else None
    },
    "timestamp": int(datetime.now().timestamp() * 1000)
}
try:
    with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
        f.write(json.dumps(log_data) + "\n")
except:
    pass
# #endregion

logger = logging.getLogger(__name__)


@dataclass
class SeedingStatus:
    """Status tracking for seeding operations."""
    stage: str  # users, templates, deals, documents, etc.
    progress: float = 0.0  # 0.0 to 1.0
    total: int = 0
    current: int = 0
    errors: List[str] = field(default_factory=list)
    status: str = "pending"  # pending, running, completed, failed
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None


# Global seeding status store - persists across requests
# This is necessary because each API request creates a new DemoDataService instance
_global_seeding_status: Dict[str, SeedingStatus] = {}


class DemoDataService:
    """Service for seeding and managing demo data."""
    
    def __init__(self, db: Session):
        """
        Initialize demo data service.
        
        Args:
            db: Database session
        """
        self.db = db
        # Use global status store to persist status across requests
        self._status = _global_seeding_status
        self._progress_callback: Optional[Callable[[str, SeedingStatus], None]] = None
        # Initialize file storage service for creating demo files
        self.file_storage = FileStorageService(base_storage_path="storage/deals")
    
    def set_progress_callback(self, callback: Callable[[str, SeedingStatus], None]):
        """Set callback for progress updates."""
        self._progress_callback = callback
    
    def _create_synthetic_document_file(
        self,
        document: Document,
        deal: Deal,
        file_type: str = "pdf"
    ) -> Optional[str]:
        """
        Create a synthetic document file (PDF or DOCX) for demo purposes.
        
        Args:
            document: Document object
            deal: Deal object
            file_type: File type ("pdf" or "docx")
            
        Returns:
            Path to created file, or None if creation failed
        """
        try:
            # Generate filename
            safe_title = "".join(c for c in document.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{safe_title.replace(' ', '_')}_{document.id}.{file_type}"
            
            # Create deal folder if needed
            deal_folder_path = self.file_storage.create_deal_folder(
                user_id=deal.applicant_id,
                deal_id=deal.deal_id
            )
            
            # Generate file content
            if file_type == "docx":
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = DocxDocument()
                    
                    # Title
                    title = doc.add_heading(document.title, 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Document metadata
                    doc.add_paragraph(f"Borrower: {document.borrower_name or 'N/A'}")
                    doc.add_paragraph(f"LEI: {document.borrower_lei or 'N/A'}")
                    doc.add_paragraph(f"Governing Law: {document.governing_law or 'N/A'}")
                    if document.total_commitment:
                        doc.add_paragraph(f"Total Commitment: {document.total_commitment:,.2f} {document.currency or 'USD'}")
                    if document.agreement_date:
                        doc.add_paragraph(f"Agreement Date: {document.agreement_date}")
                    if document.sustainability_linked:
                        doc.add_paragraph("Sustainability-Linked Loan: Yes")
                    
                    # Add content sections
                    doc.add_heading("FACILITY DETAILS", 1)
                    doc.add_paragraph("This is a synthetic demo document generated for testing purposes.")
                    doc.add_paragraph("The document contains placeholder content representing a credit agreement.")
                    
                    doc.add_heading("TERMS AND CONDITIONS", 1)
                    doc.add_paragraph("Terms and conditions would be detailed here in a real document.")
                    
                    # Save to bytes
                    import io
                    file_content = io.BytesIO()
                    doc.save(file_content)
                    file_content.seek(0)
                    content_bytes = file_content.read()
                    
                except ImportError:
                    logger.warning("python-docx not available, creating text file instead")
                    # Fallback to text file
                    content_bytes = f"""
{document.title}

Borrower: {document.borrower_name or 'N/A'}
LEI: {document.borrower_lei or 'N/A'}
Governing Law: {document.governing_law or 'N/A'}
Total Commitment: {document.total_commitment:,.2f if document.total_commitment else 'N/A'} {document.currency or 'USD'}
Agreement Date: {document.agreement_date or 'N/A'}

This is a synthetic demo document generated for testing purposes.
""".encode('utf-8')
                    filename = filename.replace('.docx', '.txt')
            else:  # PDF or fallback
                # Create a simple text-based "PDF" representation
                # In production, use a proper PDF library like reportlab or fpdf
                content_bytes = f"""
{document.title}

Borrower: {document.borrower_name or 'N/A'}
LEI: {document.borrower_lei or 'N/A'}
Governing Law: {document.governing_law or 'N/A'}
Total Commitment: {document.total_commitment:,.2f if document.total_commitment else 'N/A'} {document.currency or 'USD'}
Agreement Date: {document.agreement_date or 'N/A'}

This is a synthetic demo document generated for testing purposes.
The document contains placeholder content representing a credit agreement.
""".encode('utf-8')
                filename = filename.replace('.pdf', '.txt')
            
            # Store file using file storage service
            file_path = self.file_storage.store_deal_document(
                user_id=deal.applicant_id,
                deal_id=deal.deal_id,
                document_id=document.id,
                filename=filename,
                content=content_bytes,
                subdirectory="documents"
            )
            
            logger.debug(f"Created synthetic document file: {file_path}")
            return file_path
            
        except Exception as e:
            logger.warning(f"Failed to create synthetic document file for document {document.id}: {e}")
            return None
    
    def _update_status(self, stage: str, **kwargs):
        """Update status for a stage."""
        if stage not in self._status:
            self._status[stage] = SeedingStatus(stage=stage)
        
        status = self._status[stage]
        for key, value in kwargs.items():
            if hasattr(status, key):
                setattr(status, key, value)
        
        if self._progress_callback:
            self._progress_callback(stage, status)
    
    def seed_all(
        self,
        seed_users: bool = True,
        seed_templates: bool = True,
        seed_policies: bool = True,
        seed_policy_templates: bool = True,
        seed_securitization: bool = False,
        dry_run: bool = False
    ) -> Dict[str, Any]:
        """
        Seed all demo data.
        
        Args:
            seed_users: Whether to seed users
            seed_templates: Whether to seed templates
            seed_policies: Whether to seed policies
            seed_policy_templates: Whether to seed policy templates
            seed_securitization: Whether to seed securitization pools (requires deals to exist)
            dry_run: If True, preview what would be seeded without committing
            
        Returns:
            Dictionary with seeding results
        """
        results = {
            "users": {"created": 0, "updated": 0, "errors": []},
            "templates": {"created": 0, "updated": 0, "errors": []},
            "policies": {"created": 0, "updated": 0, "errors": []},
            "policy_templates": {"created": 0, "updated": 0, "errors": []},
            "securitization": {"created": 0, "errors": []},
        }
        
        try:
            if seed_users:
                results["users"] = self.seed_users(dry_run=dry_run)
            
            if seed_templates:
                results["templates"] = self.seed_templates(dry_run=dry_run)
            
            if seed_policies:
                results["policies"] = self.seed_policies(dry_run=dry_run)
            
            if seed_policy_templates:
                results["policy_templates"] = self.seed_policy_templates(dry_run=dry_run)
            
            # Securitization requires deals and loan assets to exist
            if seed_securitization:
                try:
                    deals = self.db.query(Deal).filter(Deal.is_demo == True).all()
                    loan_assets = self.db.query(LoanAsset).all()
                    if deals or loan_assets:
                        pools = self.create_demo_securitization_pools(deals, loan_assets, num_pools=3)
                        results["securitization"] = {"created": len(pools), "errors": []}
                    else:
                        results["securitization"] = {"created": 0, "errors": ["No deals or loan assets available for securitization"]}
                except Exception as e:
                    results["securitization"] = {"created": 0, "errors": [str(e)]}
            
            if not dry_run:
                self.db.commit()
            
            return results
        except Exception as e:
            logger.error(f"Error seeding demo data: {e}", exc_info=True)
            if not dry_run:
                self.db.rollback()
            raise
    
    def seed_users(self, force: bool = False, dry_run: bool = False) -> Dict[str, Any]:
        """
        Seed demo users.
        
        Args:
            force: If True, update existing users
            dry_run: If True, preview without committing
            
        Returns:
            Dictionary with created/updated counts, errors, and user credentials
        """
        self._update_status("users", status="running", started_at=datetime.utcnow())
        
        try:
            # Import seed function and user definitions
            from scripts.seed_demo_users import seed_demo_users, DEMO_USERS
            
            # Call seed function with seed_all_roles=True to ensure all users are created
            # (API calls should seed all roles by default, not rely on environment variables)
            count = seed_demo_users(self.db, force=force, seed_all_roles=True)
            
            # Validate and migrate profile data against UserProfileData schema for all users
            if not dry_run:
                users = self.db.query(User).filter(
                    User.email.in_([u["email"] for u in DEMO_USERS])
                ).all()
                
                for user in users:
                    if user.profile_data:
                        try:
                            # Validate profile_data against UserProfileData schema
                            UserProfileData.model_validate(user.profile_data)
                        except Exception as e:
                            # Check if it's an old format issue (company as string)
                            profile_data = user.profile_data
                            if isinstance(profile_data.get("company"), str):
                                # Migrate old format: convert company string to CompanyInfo structure
                                logger.info(f"Migrating old profile_data format for {user.email}")
                                old_company_name = profile_data.get("company", "")
                                
                                # Generate new comprehensive profile data
                                from scripts.seed_demo_users import _generate_comprehensive_profile_data
                                role = UserRole(user.role)
                                new_profile_data = _generate_comprehensive_profile_data(role, user.email)
                                
                                # Update user with new profile data
                                user.profile_data = new_profile_data
                                self.db.commit()
                                logger.info(f"Successfully migrated profile_data for {user.email}")
                            else:
                                logger.warning(f"Profile data validation warning for {user.email}: {e}")
                                # Don't fail, just log warning - profile_data may have extra fields
            
            # Prepare user credentials list
            user_credentials = []
            for user_data in DEMO_USERS:
                user_credentials.append({
                    "email": user_data["email"],
                    "password": user_data["password"],
                    "role": user_data["role"].value,
                    "display_name": user_data["display_name"]
                })
            
            result = {
                "created": count if not force else 0,
                "updated": count if force else 0,
                "errors": [],
                "user_credentials": user_credentials
            }
            
            self._update_status("users", status="completed", completed_at=datetime.utcnow(), current=count, total=count, progress=1.0)
            return result
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error seeding users: {error_msg}", exc_info=True)
            self._update_status("users", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
            return {"created": 0, "updated": 0, "errors": [error_msg], "user_credentials": []}
    
    def seed_templates(self, dry_run: bool = False) -> Dict[str, Any]:
        """
        Seed templates.
        
        Args:
            dry_run: If True, preview without committing
            
        Returns:
            Dictionary with created/updated counts and errors
        """
        self._update_status("templates", status="running", started_at=datetime.utcnow())
        
        try:
            # Import seed function
            from scripts.seed_templates import seed_templates
            import json
            from pathlib import Path
            
            # Load templates metadata
            templates_file = Path("data/templates_metadata.json")
            if not templates_file.exists():
                raise FileNotFoundError(f"Templates metadata file not found: {templates_file}")
            
            with open(templates_file, "r", encoding="utf-8") as f:
                templates_data = json.load(f)
            
            # Audit: Check all templates from metadata are present
            expected_template_codes = {t["template_code"] for t in templates_data}
            logger.info(f"Audit: Found {len(expected_template_codes)} templates in metadata file")
            
            # Call seed function
            count = seed_templates(self.db, templates_data)
            
            # Verify seeded templates
            from app.db.models import LMATemplate
            seeded_templates = self.db.query(LMATemplate).all()
            seeded_template_codes = {t.template_code for t in seeded_templates}
            
            missing_templates = expected_template_codes - seeded_template_codes
            if missing_templates:
                logger.warning(f"Audit: {len(missing_templates)} templates from metadata not found in database: {missing_templates}")
            else:
                logger.info(f"Audit: All {len(expected_template_codes)} templates from metadata are seeded")
            
            result = {
                "created": count,
                "updated": 0,
                "errors": [],
                "audit": {
                    "expected_count": len(expected_template_codes),
                    "seeded_count": len(seeded_template_codes),
                    "missing_templates": list(missing_templates) if missing_templates else []
                }
            }
            
            self._update_status("templates", status="completed", completed_at=datetime.utcnow(), current=count, total=count, progress=1.0)
            return result
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error seeding templates: {error_msg}", exc_info=True)
            self._update_status("templates", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
            return {"created": 0, "updated": 0, "errors": [error_msg]}
    
    def seed_policies(self, dry_run: bool = False) -> Dict[str, Any]:
        """
        Seed policies from YAML files.
        
        Args:
            dry_run: If True, preview without committing
            
        Returns:
            Dictionary with created/updated counts and errors
        """
        self._update_status("policies", status="running", started_at=datetime.utcnow())
        
        try:
            # Import seed function
            from scripts.seed_policies import seed_policies_from_yaml
            from app.db.models import User
            from pathlib import Path
            
            # Audit: Check all YAML files in policies directory
            policies_dir = Path("app/policies")
            yaml_files = list(policies_dir.glob("**/*.yaml")) + list(policies_dir.glob("**/*.yml"))
            expected_policy_files = {f.name for f in yaml_files}
            logger.info(f"Audit: Found {len(expected_policy_files)} YAML policy files: {sorted([f.name for f in yaml_files])}")
            
            # Get admin user ID
            admin = self.db.query(User).filter(User.role == 'admin').first()
            admin_user_id = admin.id if admin else 1
            
            # Call seed function
            count = seed_policies_from_yaml(self.db, admin_user_id)
            
            # Verify seeded policies (check if policies table has entries)
            from app.db.models import Policy
            seeded_policies = self.db.query(Policy).all()
            logger.info(f"Audit: {len(seeded_policies)} policies seeded in database")
            
            if len(seeded_policies) < len(expected_policy_files):
                logger.warning(f"Audit: Policy count mismatch - {len(expected_policy_files)} YAML files but {len(seeded_policies)} policies in database")
            
            result = {
                "created": count,
                "updated": 0,
                "errors": [],
                "audit": {
                    "expected_yaml_files": len(expected_policy_files),
                    "seeded_policies": len(seeded_policies),
                    "yaml_files": sorted([f.name for f in yaml_files])
                }
            }
            
            self._update_status("policies", status="completed", completed_at=datetime.utcnow(), current=count, total=count, progress=1.0)
            return result
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error seeding policies: {error_msg}", exc_info=True)
            self._update_status("policies", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
            return {"created": 0, "updated": 0, "errors": [error_msg]}
    
    def seed_policy_templates(self, dry_run: bool = False) -> Dict[str, Any]:
        """
        Seed policy templates.
        
        Args:
            dry_run: If True, preview without committing
            
        Returns:
            Dictionary with created/updated counts and errors
        """
        self._update_status("policy_templates", status="running", started_at=datetime.utcnow())
        
        try:
            # Import seed function (if exists)
            try:
                from scripts.seed_policy_templates import seed_policy_templates
                count = seed_policy_templates(self.db)
            except ImportError:
                logger.warning("seed_policy_templates not available, skipping")
                count = 0
            
            result = {
                "created": count,
                "updated": 0,
                "errors": []
            }
            
            self._update_status("policy_templates", status="completed", completed_at=datetime.utcnow(), current=count, total=count, progress=1.0)
            return result
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error seeding policy templates: {error_msg}", exc_info=True)
            self._update_status("policy_templates", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
            return {"created": 0, "updated": 0, "errors": [error_msg]}
    
    def get_seeding_status(self, stage: Optional[str] = None) -> Dict[str, Any]:
        """
        Get current seeding status.
        
        Args:
            stage: Optional stage name to get status for specific stage
            
        Returns:
            Dictionary with status information
        """
        if stage:
            status = self._status.get(stage)
            if status:
                return {
                    "stage": status.stage,
                    "progress": status.progress,
                    "total": status.total,
                    "current": status.current,
                    "errors": status.errors,
                    "status": status.status,
                    "started_at": status.started_at.isoformat() if status.started_at else None,
                    "completed_at": status.completed_at.isoformat() if status.completed_at else None,
                }
            return None
        
        # Return all statuses
        return {
            stage: {
                "stage": status.stage,
                "progress": status.progress,
                "total": status.total,
                "current": status.current,
                "errors": status.errors,
                "status": status.status,
                "started_at": status.started_at.isoformat() if status.started_at else None,
                "completed_at": status.completed_at.isoformat() if status.completed_at else None,
            }
            for stage, status in self._status.items()
        }
    
    def _generate_cdm_for_deal(
        self,
        deal_type: str,
        seed: Optional[int] = None,
        scenario: Optional[str] = None,
        use_cache: bool = True
    ) -> CreditAgreement:
        """
        Generate CDM data for a deal using AI with optional caching.
        
        Args:
            deal_type: Type of deal (loan_application, refinancing, restructuring)
            seed: Random seed for reproducibility
            scenario: Scenario template (if None, randomly selects)
            use_cache: Whether to use cache (default: True)
            
        Returns:
            CreditAgreement instance
        """
        # Check cache first if enabled
        if use_cache:
            try:
                from app.services.demo_data_cache import get_demo_cache
                cache = get_demo_cache()
                
                cached_cdm = cache.get_cached_deal(seed=seed or 0, deal_type=deal_type, scenario=scenario)
                if cached_cdm:
                    logger.debug(f"Cache hit for deal (seed={seed}, type={deal_type}, scenario={scenario})")
                    try:
                        return CreditAgreement(**cached_cdm.get("cdm", cached_cdm))
                    except Exception as e:
                        logger.warning(f"Failed to deserialize cached CDM: {e}")
                        # Fall through to generation
            except Exception as e:
                logger.warning(f"Cache lookup failed: {e}")
                # Fall through to generation
        
        # Generate new CDM
        from app.chains.deal_generation_chain import generate_cdm_for_deal
        
        cdm = generate_cdm_for_deal(
            deal_type=deal_type,
            scenario=scenario,
            seed=seed
        )
        
        # Cache the result if enabled
        if use_cache:
            try:
                from app.services.demo_data_cache import get_demo_cache
                cache = get_demo_cache()
                
                cache.cache_deal(
                    seed=seed or 0,
                    deal_type=deal_type,
                    deal_data={"cdm": cdm.model_dump()},
                    scenario=scenario
                )
                logger.debug(f"Cached deal (seed={seed}, type={deal_type}, scenario={scenario})")
            except Exception as e:
                logger.warning(f"Failed to cache deal: {e}")
        
        return cdm
    
    def generate_deal_data(
        self,
        count: int = 12,
        deal_types: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Generate deal data using AI.
        
        Args:
            count: Number of deals to generate
            deal_types: List of deal types to generate (if None, uses defaults)
            
        Returns:
            List of deal data dictionaries
        """
        if deal_types is None:
            deal_types = ["loan_application", "refinancing", "restructuring"]
        
        deals = []
        self._update_status("deals", status="running", started_at=datetime.utcnow(), total=count, current=0)
        
        for i in range(count):
            try:
                # Select deal type
                deal_type = random.choice(deal_types) if deal_types else "loan_application"
                
                # Generate CDM
                cdm = self._generate_cdm_for_deal(deal_type=deal_type, seed=i)
                
                # Convert to dict
                deal_data = {
                    "cdm": cdm.model_dump(),
                    "deal_type": deal_type,
                    "seed": i
                }
                
                deals.append(deal_data)
                self._update_status("deals", current=i + 1, progress=(i + 1) / count)
                
            except Exception as e:
                error_msg = f"Failed to generate deal {i + 1}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("deals", errors=[error_msg])
                continue
        
        self._update_status("deals", status="completed", completed_at=datetime.utcnow())
        return deals
    
    def create_demo_deals(self, count: int = 12, seed_securitization: bool = False) -> List[Deal]:
        """
        Create complete demo deals with full data flow: Applications → Deals → Documents → Workflows.
        
        Only creates new deals if the requested count exceeds existing demo deals.
        
        Args:
            count: Number of deals to have (not number to create)
            seed_securitization: Whether to create securitization pools
            
        Returns:
            List of created Deal objects (empty if count satisfied by existing)
        """
        # Check existing demo deals
        existing_demo_deals = self.db.query(Deal).filter(Deal.is_demo == True).count()
        
        # If we already have enough demo deals, skip creation
        if existing_demo_deals >= count:
            logger.info(f"Skipping deal creation: {existing_demo_deals} demo deals exist, {count} requested")
            self._update_status("deals", status="completed", completed_at=datetime.utcnow(), 
                              total=count, current=count, progress=1.0)
            return []
        
        # Only create the difference
        deals_to_create = count - existing_demo_deals
        logger.info(f"Creating {deals_to_create} new demo deals (have {existing_demo_deals}, want {count})")
        
        self._update_status("deals", status="running", started_at=datetime.utcnow(), total=deals_to_create, current=0)
        
        # Step 1: Generate Applications
        applications = self._generate_applications(deals_to_create)
        
        # Step 2: Create Deals from Applications
        deals = self._create_deals_from_applications(applications)
        
        # Step 3: Generate Documents for Deals
        documents = self._generate_deal_documents(deals)
        
        # Step 4: Create DocumentVersions for Documents
        document_versions = self._create_document_versions(documents)
        
        # Step 4b: Create revision history for 30% of documents
        document_ids_for_revision = [d.id for d in documents]
        revision_versions = self.create_document_revision_history(document_ids_for_revision)
        document_versions.extend(revision_versions)
        
        # Step 5: Generate Workflows for Documents
        workflows = self._create_workflows_for_documents(documents)
        
        # Step 6: Generate LoanAssets for sustainability-linked deals
        loan_assets = self._generate_loan_assets_for_deals(deals)
        
        # Step 6b: Verify correlation - ensure deals with files have loan assets
        # This ensures 100% correlation between deals with attached files and loan assets
        deals_with_files = []
        for deal in deals:
            try:
                deal_files = self.file_storage.get_deal_documents(
                    user_id=deal.applicant_id,
                    deal_id=deal.deal_id,
                    subdirectory="documents"
                )
                if deal_files:
                    deals_with_files.append(deal)
            except Exception:
                pass
        
        # For deals with files that are sustainability-linked, ensure they have loan assets
        sustainability_deals_with_files = [
            d for d in deals_with_files 
            if d.deal_data and d.deal_data.get("sustainability_linked", False)
        ]
        
        for deal in sustainability_deals_with_files:
            existing_loan_asset = self.db.query(LoanAsset).filter(
                LoanAsset.loan_id == deal.deal_id
            ).first()
            if not existing_loan_asset:
                logger.warning(f"Deal {deal.deal_id} has files but no loan asset - this should not happen for sustainability-linked deals")
        
        logger.info(f"File attachment correlation: {len(deals_with_files)} deals have files, {len(sustainability_deals_with_files)} are sustainability-linked with files")
        
        # #region agent log
        log_data = {
            "sessionId": "debug-session",
            "runId": "post-fix",
            "hypothesisId": "A",
            "location": "demo_data_service.py:505",
            "message": "Post-fix: Before settings access - verification",
            "data": {
                "loan_assets_count": len(loan_assets) if loan_assets else 0,
                "settings_in_globals": "settings" in globals(),
                "settings_accessible": "settings" in globals() and hasattr(settings, "ENHANCED_SATELLITE_ENABLED"),
                "enhanced_satellite_value": getattr(settings, "ENHANCED_SATELLITE_ENABLED", None) if "settings" in globals() else None
            },
            "timestamp": int(datetime.now().timestamp() * 1000)
        }
        with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
            f.write(json.dumps(log_data) + "\n")
        # #endregion
        
        # #region agent log
        try:
            test_value = settings.ENHANCED_SATELLITE_ENABLED
            log_data = {
                "sessionId": "debug-session",
                "runId": "post-fix",
                "hypothesisId": "A",
                "location": "demo_data_service.py:506",
                "message": "Post-fix: Settings access test - SUCCESS",
                "data": {
                    "settings_type": type(settings).__name__,
                    "enhanced_satellite_value": test_value,
                    "access_successful": True
                },
                "timestamp": int(datetime.now().timestamp() * 1000)
            }
        except NameError as e:
            log_data = {
                "sessionId": "debug-session",
                "runId": "post-fix",
                "hypothesisId": "A",
                "location": "demo_data_service.py:506",
                "message": "Post-fix: Settings access test - NameError (FIX FAILED)",
                "data": {"error": str(e), "error_type": type(e).__name__, "access_successful": False},
                "timestamp": int(datetime.now().timestamp() * 1000)
            }
        except Exception as e:
            log_data = {
                "sessionId": "debug-session",
                "runId": "post-fix",
                "hypothesisId": "A",
                "location": "demo_data_service.py:506",
                "message": "Post-fix: Settings access test - Other Error",
                "data": {"error": str(e), "error_type": type(e).__name__, "access_successful": False},
                "timestamp": int(datetime.now().timestamp() * 1000)
            }
        with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
            f.write(json.dumps(log_data) + "\n")
        # #endregion
        
        # Step 6b: Create green finance assessments for loan assets
        if loan_assets and settings.ENHANCED_SATELLITE_ENABLED:
            try:
                green_finance_assessments = self._create_green_finance_assessments(loan_assets, deals)
                logger.info(f"Created {len(green_finance_assessments)} green finance assessments for demo loan assets")
            except Exception as e:
                logger.warning(f"Failed to create green finance assessments: {e}")
        
        # Step 7: Generate documents from templates for approved deals
        approved_deals = [d for d in deals if d.status in [DealStatus.APPROVED.value, DealStatus.ACTIVE.value]]
        generated_docs = []
        for deal in approved_deals:
            try:
                docs = self.generate_documents_from_templates(deal.id)
                generated_docs.extend(docs)
            except Exception as e:
                logger.warning(f"Failed to generate documents from templates for deal {deal.id}: {e}")
        
        # Step 8: Create deal notes
        deal_notes = self.create_deal_notes([d.id for d in deals])
        
        # Step 9: Create policy decisions for documents
        policy_decisions = self.create_policy_decisions_for_documents([d.id for d in documents])
        
        # Step 9b: Create green finance policy decisions for loan assets (if enabled)
        if loan_assets and settings.ENHANCED_SATELLITE_ENABLED:
            try:
                green_policy_decisions = self._create_green_finance_policy_decisions(loan_assets, deals)
                policy_decisions.extend(green_policy_decisions)
                logger.info(f"Created {len(green_policy_decisions)} green finance policy decisions")
            except Exception as e:
                logger.warning(f"Failed to create green finance policy decisions: {e}")
        
        # Step 9c: Create securitization pools (if enabled)
        securitization_pools = []
        if seed_securitization:
            try:
                securitization_pools = self.create_demo_securitization_pools(deals, loan_assets, num_pools=3)
                logger.info(f"Created {len(securitization_pools)} securitization pools")
            except Exception as e:
                logger.warning(f"Failed to create securitization pools: {e}")
        
        # Step 10: Store generated documents as files
        self._store_generated_documents(deals, documents, document_versions, generated_docs)
        
        # Step 11: Index documents in ChromaDB
        self.index_demo_documents([d.id for d in documents])
        
        self._update_status("deals", status="completed", completed_at=datetime.utcnow(), current=len(deals), progress=1.0)
        
        return deals
    
    def _generate_applications(self, count: int) -> List[Application]:
        """
        Generate Applications (12-15) with business/individual types, status distribution.
        
        Args:
            count: Number of applications to generate
            
        Returns:
            List of Application objects
        """
        self._update_status("applications", status="running", started_at=datetime.utcnow(), total=count, current=0)
        
        # Get demo applicants
        applicants = self.db.query(User).filter(User.role == UserRole.APPLICANT.value).all()
        if not applicants:
            raise ValueError("No applicant users found. Please seed demo users first.")
        
        applications = []
        status_distribution = {
            ApplicationStatus.DRAFT.value: 0.20,
            ApplicationStatus.SUBMITTED.value: 0.30,
            ApplicationStatus.UNDER_REVIEW.value: 0.30,
            ApplicationStatus.APPROVED.value: 0.15,
            ApplicationStatus.REJECTED.value: 0.05
        }
        
        for i in range(count):
            try:
                # Select application type (80% business, 20% individual)
                app_type = ApplicationType.BUSINESS.value if random.random() < 0.80 else ApplicationType.INDIVIDUAL.value
                
                # Select status based on distribution
                status = random.choices(
                    list(status_distribution.keys()),
                    weights=list(status_distribution.values())
                )[0]
                
                # Select applicant
                applicant = random.choice(applicants)
                
                # Select industry using weights
                from app.prompts.demo.deal_generation import get_industry_weights, get_industry_config
                industry_weights = get_industry_weights()
                industries = list(industry_weights.keys())
                weights = list(industry_weights.values())
                industry = random.choices(industries, weights=weights)[0]
                
                # Get industry-specific config
                industry_config = get_industry_config(industry)
                
                # Generate application data using industry config
                loan_amount_min, loan_amount_max = industry_config["loan_amount_range"]
                loan_amount = random.randint(loan_amount_min, loan_amount_max)
                years_in_business = random.randint(2, 50) if app_type == ApplicationType.BUSINESS.value else random.randint(1, 30)
                annual_revenue = loan_amount * random.uniform(2, 10)
                credit_score = random.randint(650, 850)
                
                application_data = {
                    "loan_amount": loan_amount,
                    "purpose": random.choice(["Working capital", "Expansion", "Refinancing", "Equipment purchase", "Inventory"]),
                    "industry": industry,
                    "years_in_business": years_in_business,
                    "annual_revenue": annual_revenue,
                    "credit_score": credit_score
                }
                
                business_data = None
                if app_type == ApplicationType.BUSINESS.value:
                    business_data = {
                        "company_name": f"Demo Company {i + 1}",
                        "tax_id": f"{random.randint(10, 99)}-{random.randint(1000000, 9999999)}",
                        "legal_structure": random.choice(["Corporation", "LLC", "Partnership"]),
                        "number_of_employees": random.randint(10, 5000),
                        "business_address": f"{random.randint(1, 9999)} Business St, City, ST {random.randint(10000, 99999)}"
                    }
                
                # Generate timestamps
                base_date = datetime.utcnow() - timedelta(days=random.randint(30, 180))
                submitted_at = base_date if status != ApplicationStatus.DRAFT.value else None
                reviewed_at = None
                approved_at = None
                rejected_at = None
                
                if status in [ApplicationStatus.UNDER_REVIEW.value, ApplicationStatus.APPROVED.value, ApplicationStatus.REJECTED.value]:
                    reviewed_at = submitted_at + timedelta(days=random.randint(1, 7)) if submitted_at else None
                
                if status == ApplicationStatus.APPROVED.value:
                    approved_at = reviewed_at + timedelta(days=random.randint(1, 14)) if reviewed_at else None
                
                if status == ApplicationStatus.REJECTED.value:
                    rejected_at = reviewed_at + timedelta(days=random.randint(1, 3)) if reviewed_at else None
                    rejection_reason = random.choice([
                        "Insufficient credit history",
                        "Inadequate collateral",
                        "High debt-to-income ratio",
                        "Unstable revenue stream",
                        "Industry risk concerns",
                        "Incomplete documentation"
                    ])
                else:
                    rejection_reason = None
                
                # Generate individual_data for individual applications
                individual_data = None
                if app_type == ApplicationType.INDIVIDUAL.value:
                    individual_data = {
                        "first_name": f"John",
                        "last_name": f"Doe{i + 1}",
                        "date_of_birth": (datetime.utcnow() - timedelta(days=random.randint(25*365, 65*365))).date().isoformat(),
                        "ssn_last_4": f"{random.randint(1000, 9999)}",
                        "employment_status": random.choice(["Employed", "Self-employed", "Retired"]),
                        "annual_income": annual_revenue,
                        "residential_address": f"{random.randint(1, 9999)} Residential Ave, City, ST {random.randint(10000, 99999)}",
                        "years_at_address": random.randint(1, 20)
                    }
                
                # Create application
                application = Application(
                    application_type=app_type,
                    status=status,
                    user_id=applicant.id,
                    submitted_at=submitted_at,
                    reviewed_at=reviewed_at,
                    approved_at=approved_at,
                    rejected_at=rejected_at,
                    rejection_reason=rejection_reason,
                    application_data=application_data,
                    business_data=business_data,
                    individual_data=individual_data
                )
                
                self.db.add(application)
                applications.append(application)
                self._update_status("applications", current=i + 1, progress=(i + 1) / count)
                
            except Exception as e:
                error_msg = f"Failed to generate application {i + 1}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("applications", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("applications", status="completed", completed_at=datetime.utcnow())
        
        return applications
    
    def _create_deals_from_applications(self, applications: List[Application]) -> List[Deal]:
        """
        Create Deals from Applications with deal_id format, status mapping, deal_data JSONB.
        
        Args:
            applications: List of Application objects
            
        Returns:
            List of Deal objects
        """
        self._update_status("deals_creation", status="running", started_at=datetime.utcnow(), total=len(applications), current=0)
        
        deals = []
        now = datetime.utcnow()
        deal_id_prefix = f"DEAL-{now.year}-{now.month:02d}-"
        
        # Find the highest existing deal counter for this month to avoid conflicts
        existing_deals = self.db.query(Deal).filter(
            Deal.deal_id.like(f"{deal_id_prefix}%")
        ).all()
        
        max_counter = 0
        if existing_deals:
            for deal in existing_deals:
                try:
                    # Extract counter from deal_id (e.g., "DEAL-2026-01-001" -> 1)
                    counter_str = deal.deal_id.split("-")[-1]
                    counter = int(counter_str)
                    max_counter = max(max_counter, counter)
                except (ValueError, IndexError):
                    continue
        
        deal_counter = max_counter + 1
        
        for application in applications:
            try:
                # Only create deals for submitted+ applications
                if application.status == ApplicationStatus.DRAFT.value:
                    continue
                
                # Generate deal_id
                deal_id = f"{deal_id_prefix}{deal_counter:03d}"
                
                # Check if deal_id already exists, delete it and related data if it does
                existing_deal = self.db.query(Deal).filter(Deal.deal_id == deal_id).first()
                if existing_deal:
                    logger.info(f"Deleting existing demo deal {deal_id} and related data to overwrite with new data")
                    # Delete related data
                    self.db.query(Document).filter(Document.deal_id == existing_deal.id).delete()
                    self.db.query(DealNote).filter(DealNote.deal_id == existing_deal.id).delete()
                    # Delete workflows via documents (cascade should handle this, but be explicit)
                    doc_ids = [d.id for d in self.db.query(Document).filter(Document.deal_id == existing_deal.id).all()]
                    if doc_ids:
                        self.db.query(Workflow).filter(Workflow.document_id.in_(doc_ids)).delete()
                    # Delete the deal
                    self.db.delete(existing_deal)
                    self.db.flush()  # Flush to ensure deletion before creating new deal
                
                deal_counter += 1
                
                # Assign realistic deal status based on application status and deal age
                deal_status = self._assign_realistic_deal_status(application, deal_counter)
                
                # Get deal data from application
                loan_amount = application.application_data.get("loan_amount", 1000000) if application.application_data else 1000000
                industry = application.application_data.get("industry", "Technology") if application.application_data else "Technology"
                
                # Get industry-specific config
                from app.prompts.demo.deal_generation import get_industry_config
                industry_config = get_industry_config(industry)
                
                # Use industry-specific ranges
                term_min, term_max = industry_config["term_range"]
                term_years = random.randint(term_min, term_max)
                
                rate_min, rate_max = industry_config["interest_rate_range"]
                interest_rate = random.uniform(rate_min, rate_max)
                
                sustainability_linked = random.random() < 0.30  # 30% sustainability-linked
                
                # Enhanced deal_data with comprehensive metadata
                deal_data = {
                    "loan_amount": loan_amount,
                    "interest_rate": interest_rate,
                    "term_years": term_years,
                    "industry": industry,
                    "collateral_type": random.choice(industry_config["collateral_types"]),
                    "sustainability_linked": sustainability_linked,
                    "is_demo": True,  # Explicitly mark as demo
                    "created_via": "demo_data_service",
                    "loan_purpose": application.application_data.get("purpose", "Working capital") if application.application_data else "Working capital",
                    "origination_date": (application.submitted_at or application.created_at).isoformat() if (application.submitted_at or application.created_at) else None,
                    "maturity_date": None,  # Will be calculated based on term_years if needed
                    "amortization_schedule": random.choice(["Bullet", "Amortizing", "Balloon"]),
                    "prepayment_penalty": random.choice([True, False]),
                    "covenants": {
                        "debt_service_coverage_ratio": random.uniform(1.2, 2.0),
                        "loan_to_value_ratio": random.uniform(0.5, 0.85),
                        "debt_to_equity_ratio": random.uniform(0.3, 2.0)
                    },
                    "risk_rating": random.choice(["A", "BBB", "BB", "B", "CCC"]),
                    "syndication_status": random.choice(["Sole lender", "Syndicated", "Club deal"]),
                    "currency": "USD"
                }
                
                # Calculate maturity date if term_years is set
                if deal_data.get("term_years") and (application.submitted_at or application.created_at):
                    base_date = application.submitted_at or application.created_at
                    maturity_date = base_date + timedelta(days=deal_data["term_years"] * 365)
                    deal_data["maturity_date"] = maturity_date.isoformat()
                
                if sustainability_linked:
                    deal_data["esg_targets"] = {
                        "ndvi_threshold": random.uniform(0.70, 0.85),
                        "verification_frequency": "Quarterly",
                        "carbon_reduction_target": random.uniform(10, 30),  # Percentage
                        "renewable_energy_target": random.uniform(20, 50),  # Percentage
                        "water_conservation_target": random.uniform(15, 40)  # Percentage
                    }
                    deal_data["esg_penalties"] = {
                        "penalty_bps": random.randint(25, 50),  # Basis points
                        "penalty_applicable": True
                    }
                
                # Create folder path
                folder_path = f"storage/deals/demo/{deal_id}/"
                
                # Set verification and notarization requirements based on deal status and type
                # Sustainability-linked deals and approved deals are more likely to require verification
                verification_required = sustainability_linked or (deal_status in [DealStatus.APPROVED.value, DealStatus.ACTIVE.value])
                verification_completed_at = None
                if verification_required and deal_status in [DealStatus.APPROVED.value, DealStatus.ACTIVE.value]:
                    # 70% of verification-required deals have completed verification
                    if random.random() < 0.70:
                        verification_completed_at = (application.approved_at or application.reviewed_at or application.submitted_at or application.created_at) + timedelta(days=random.randint(1, 14))
                
                # Notarization is required for larger deals or when deal is active/closed
                notarization_required = loan_amount > 5000000 or deal_status in [DealStatus.ACTIVE.value, DealStatus.CLOSED.value]
                notarization_completed_at = None
                if notarization_required and deal_status in [DealStatus.ACTIVE.value, DealStatus.CLOSED.value]:
                    # 60% of notarization-required deals have completed notarization
                    if random.random() < 0.60:
                        base_date = verification_completed_at or (application.approved_at or application.reviewed_at or application.submitted_at or application.created_at)
                        notarization_completed_at = base_date + timedelta(days=random.randint(1, 30))
                
                # Create deal with is_demo flag
                deal = Deal(
                    deal_id=deal_id,
                    applicant_id=application.user_id,
                    application_id=application.id,
                    status=deal_status,
                    deal_type=DealType.LOAN_APPLICATION.value,
                    is_demo=True,  # Mark as demo deal
                    deal_data=deal_data,
                    folder_path=folder_path,
                    verification_required=verification_required,
                    verification_completed_at=verification_completed_at,
                    notarization_required=notarization_required,
                    notarization_completed_at=notarization_completed_at,
                    created_at=application.submitted_at or application.created_at,
                    updated_at=application.approved_at or application.reviewed_at or application.submitted_at or application.created_at
                )
                
                self.db.add(deal)
                deals.append(deal)
                
                self._update_status("deals_creation", current=len(deals), progress=len(deals) / len(applications))
                
            except Exception as e:
                error_msg = f"Failed to create deal from application {application.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("deals_creation", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("deals_creation", status="completed", completed_at=datetime.utcnow())
        
        return deals
    
    def _assign_realistic_deal_status(self, application: Application, deal_index: int) -> str:
        """
        Assign realistic deal status based on application status and deal progression.
        
        Status distribution:
        - draft: 5%
        - submitted: 15%
        - under_review: 25%
        - approved: 20%
        - active: 15%
        - closed: 10%
        - rejected: 5%
        - restructuring: 3%
        - withdrawn: 2%
        
        Args:
            application: Application object
            deal_index: Index of deal in sequence (for variety)
            
        Returns:
            Deal status string
        """
        # Calculate deal age in days
        deal_age_days = (datetime.utcnow() - (application.created_at or datetime.utcnow())).days
        
        # Base status mapping from application status
        # Note: ApplicationStatus doesn't have PENDING, only DRAFT, SUBMITTED, UNDER_REVIEW, APPROVED, REJECTED
        base_status_mapping = {
            ApplicationStatus.DRAFT.value: DealStatus.DRAFT.value,
            ApplicationStatus.SUBMITTED.value: DealStatus.SUBMITTED.value,
            ApplicationStatus.UNDER_REVIEW.value: DealStatus.UNDER_REVIEW.value,
            ApplicationStatus.APPROVED.value: DealStatus.APPROVED.value,
            ApplicationStatus.REJECTED.value: DealStatus.REJECTED.value
        }
        
        base_status = base_status_mapping.get(application.status, DealStatus.SUBMITTED.value)
        
        # Use weighted random selection for realistic distribution
        # This ensures we get the target percentages across all deals
        rand = random.random()
        
        # Status distribution with cumulative probabilities
        if rand < 0.05:
            return DealStatus.DRAFT.value
        elif rand < 0.20:  # 5% + 15%
            return DealStatus.SUBMITTED.value
        elif rand < 0.45:  # 5% + 15% + 25%
            return DealStatus.UNDER_REVIEW.value
        elif rand < 0.65:  # 5% + 15% + 25% + 20%
            return DealStatus.APPROVED.value
        elif rand < 0.80:  # 5% + 15% + 25% + 20% + 15%
            # Active deals should have approved base status and be older
            if base_status == DealStatus.APPROVED.value and deal_age_days > 30:
                return DealStatus.ACTIVE.value
            else:
                return DealStatus.APPROVED.value
        elif rand < 0.90:  # 5% + 15% + 25% + 20% + 15% + 10%
            # Closed deals should be older active deals
            if deal_age_days > 180:
                return DealStatus.CLOSED.value
            else:
                return DealStatus.ACTIVE.value
        elif rand < 0.95:  # 5% + 15% + 25% + 20% + 15% + 10% + 5%
            return DealStatus.REJECTED.value
        elif rand < 0.98:  # 5% + 15% + 25% + 20% + 15% + 10% + 5% + 3%
            # Restructuring deals should be active deals with issues
            if base_status == DealStatus.ACTIVE.value:
                return DealStatus.RESTRUCTURING.value
            else:
                return DealStatus.APPROVED.value
        else:  # 5% + 15% + 25% + 20% + 15% + 10% + 5% + 3% + 2%
            return DealStatus.WITHDRAWN.value
    
    def _flatten_dict(self, d: Dict[str, Any], parent_key: str = '', sep: str = '.') -> Dict[str, Any]:
        """
        Recursively flatten a nested dictionary.
        
        Args:
            d: Dictionary to flatten
            parent_key: Parent key for nested dictionaries
            sep: Separator for nested keys
            
        Returns:
            Flattened dictionary
        """
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key, sep=sep).items())
            elif isinstance(v, list):
                # Handle lists by checking each item
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        items.extend(self._flatten_dict(item, f"{new_key}[{i}]", sep=sep).items())
                    else:
                        items.append((f"{new_key}[{i}]", item))
            else:
                items.append((new_key, v))
        return dict(items)
    
    def _generate_deal_documents(self, deals: List[Deal]) -> List[Document]:
        """
        Generate deal documents (18-36) with title, borrower_name, borrower_lei, etc.
        
        Args:
            deals: List of Deal objects
            
        Returns:
            List of Document objects
        """
        self._update_status("documents", status="running", started_at=datetime.utcnow(), total=len(deals) * 2, current=0)
        
        documents = []
        # Get demo users for document uploaders
        uploaders = self.db.query(User).filter(
            User.role.in_([UserRole.BANKER.value, UserRole.LAW_OFFICER.value, UserRole.ACCOUNTANT.value])
        ).all()
        
        if not uploaders:
            uploaders = self.db.query(User).limit(3).all()
        
        for deal in deals:
            try:
                # Check if documents already exist for this deal, skip if they do
                existing_docs = self.db.query(Document).filter(Document.deal_id == deal.id).all()
                if existing_docs:
                    logger.debug(f"Documents already exist for deal {deal.deal_id}, skipping document generation")
                    documents.extend(existing_docs)
                    continue
                
                # Generate 1-3 documents per deal
                doc_count = random.randint(1, 3)
                
                for doc_idx in range(doc_count):
                    # Generate CDM for this document
                    cdm = self._generate_cdm_for_deal(
                        deal_type=deal.deal_type,
                        seed=deal.id + doc_idx
                    )
                    
                    # Extract borrower info from CDM
                    borrower = next((p for p in cdm.parties if p.role == "Borrower"), None) if cdm.parties else None
                    borrower_name = borrower.name if borrower else f"Borrower {deal.deal_id}"
                    borrower_lei = borrower.lei if borrower else None
                    
                    # Calculate total commitment
                    total_commitment = Decimal(0)
                    currency = "USD"
                    if cdm.facilities:
                        for facility in cdm.facilities:
                            if facility.commitment_amount:
                                total_commitment += facility.commitment_amount.amount
                                currency = facility.commitment_amount.currency.value if hasattr(facility.commitment_amount.currency, 'value') else str(facility.commitment_amount.currency)
                    
                    # Document title
                    doc_titles = [
                        f"Credit Agreement - {borrower_name}",
                        f"Term Sheet - {deal.deal_id}",
                        f"Supporting Documents - {deal.deal_id}"
                    ]
                    title = doc_titles[doc_idx] if doc_idx < len(doc_titles) else f"Document {doc_idx + 1} - {deal.deal_id}"
                    
                    # ESG metadata
                    esg_metadata = None
                    sustainability_linked = deal.deal_data.get("sustainability_linked", False) if deal.deal_data else False
                    if sustainability_linked:
                        esg_metadata = {
                            "esg_score": random.uniform(0.65, 0.95),
                            "kpi_targets": ["NDVI > 0.75", "Carbon reduction 20%"],
                            "verification_frequency": random.choice(["Quarterly", "Annual"])
                        }
                    
                    # Create document
                    # Note: is_demo is tracked via deal.deal_data["is_demo"], not a separate field
                    # Set is_generated based on document type (term sheets are often generated)
                    is_generated = "Term Sheet" in title or random.random() < 0.15  # 15% chance for other docs
                    
                    # Set template_id and source_cdm_data for generated documents
                    template_id = None
                    source_cdm_data = None
                    if is_generated:
                        # For generated documents, try to find a matching template
                        from app.db.models import LMATemplate
                        # Try to find a template matching the document type
                        if "Term Sheet" in title:
                            template = self.db.query(LMATemplate).filter(
                                LMATemplate.name.ilike("%term sheet%")
                            ).first()
                        elif "Credit Agreement" in title:
                            template = self.db.query(LMATemplate).filter(
                                LMATemplate.category == "Credit Agreement"
                            ).first()
                        else:
                            # Get any available template
                            template = self.db.query(LMATemplate).first()
                        
                        if template:
                            template_id = template.id
                            # Store the CDM data used for generation
                            # Use model_dump_json() and parse to ensure dates are serialized as strings for JSONB storage
                            # #region agent log
                            log_data = {
                                "sessionId": "debug-session",
                                "runId": "fix-date-serialization",
                                "hypothesisId": "A",
                                "location": "demo_data_service.py:1376",
                                "message": "Serializing CDM data for JSONB storage",
                                "data": {
                                    "has_model_dump_json": hasattr(cdm, 'model_dump_json'),
                                    "has_model_dump": hasattr(cdm, 'model_dump'),
                                    "has_dict": hasattr(cdm, 'dict'),
                                    "cdm_type": type(cdm).__name__
                                },
                                "timestamp": int(datetime.now().timestamp() * 1000)
                            }
                            try:
                                with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
                                    f.write(json.dumps(log_data) + "\n")
                            except Exception:
                                pass
                            # #endregion
                            
                            if hasattr(cdm, 'model_dump_json'):
                                cdm_json = cdm.model_dump_json()
                                source_cdm_data = json.loads(cdm_json)
                                # #region agent log
                                log_data = {
                                    "sessionId": "debug-session",
                                    "runId": "fix-date-serialization",
                                    "hypothesisId": "A",
                                    "location": "demo_data_service.py:1395",
                                    "message": "Used model_dump_json() for serialization",
                                    "data": {
                                        "cdm_json_length": len(cdm_json),
                                        "source_cdm_data_type": type(source_cdm_data).__name__,
                                        "has_date_objects": any(isinstance(v, (date, datetime)) for v in self._flatten_dict(source_cdm_data).values()) if isinstance(source_cdm_data, dict) else False
                                    },
                                    "timestamp": int(datetime.now().timestamp() * 1000)
                                }
                                try:
                                    with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
                                        f.write(json.dumps(log_data) + "\n")
                                except Exception:
                                    pass
                                # #endregion
                            elif hasattr(cdm, 'model_dump'):
                                source_cdm_data = cdm.model_dump(mode='json')
                            elif hasattr(cdm, 'dict'):
                                # For Pydantic v1, convert date objects to ISO strings manually
                                source_cdm_data = cdm.dict()
                                
                                def json_serial(obj):
                                    """JSON serializer for objects not serializable by default json code"""
                                    if isinstance(obj, (date, datetime)):
                                        return obj.isoformat()
                                    raise TypeError(f"Type {type(obj)} not serializable")
                                
                                source_cdm_data = json.loads(json.dumps(source_cdm_data, default=json_serial))
                            else:
                                source_cdm_data = None
                            
                            # #region agent log
                            log_data = {
                                "sessionId": "debug-session",
                                "runId": "fix-date-serialization",
                                "hypothesisId": "A",
                                "location": "demo_data_service.py:1420",
                                "message": "Final source_cdm_data check before DB insert",
                                "data": {
                                    "source_cdm_data_is_none": source_cdm_data is None,
                                    "source_cdm_data_type": type(source_cdm_data).__name__ if source_cdm_data is not None else None,
                                    "can_serialize_to_json": False
                                },
                                "timestamp": int(datetime.now().timestamp() * 1000)
                            }
                            if source_cdm_data is not None:
                                try:
                                    json.dumps(source_cdm_data)  # Test if it's JSON serializable
                                    log_data["data"]["can_serialize_to_json"] = True
                                except TypeError as e:
                                    log_data["data"]["serialization_error"] = str(e)
                            try:
                                with open(r"c:\Users\MeMyself\creditnexus\.cursor\debug.log", "a") as f:
                                    f.write(json.dumps(log_data) + "\n")
                            except Exception:
                                pass
                            # #endregion
                    
                    document = Document(
                        title=title,
                        borrower_name=borrower_name,
                        borrower_lei=borrower_lei,
                        governing_law=cdm.governing_law or "NY",
                        total_commitment=total_commitment,
                        currency=currency,
                        agreement_date=cdm.agreement_date,
                        sustainability_linked=sustainability_linked,
                        esg_metadata=esg_metadata,
                        uploaded_by=random.choice(uploaders).id if uploaders else None,
                        deal_id=deal.id,
                        is_generated=is_generated,  # Set is_generated flag
                        template_id=template_id,  # Set template_id for generated documents
                        source_cdm_data=source_cdm_data,  # Store CDM data used for generation
                        created_at=deal.created_at + timedelta(days=random.randint(0, 2))
                    )
                    
                    self.db.add(document)
                    self.db.flush()  # Flush to get document.id
                    
                    # Attach file to 60-70% of documents (weighted towards sustainability-linked)
                    should_attach_file = False
                    if sustainability_linked:
                        # 80% of sustainability-linked documents get files
                        should_attach_file = random.random() < 0.80
                    else:
                        # 60% of regular documents get files
                        should_attach_file = random.random() < 0.60
                    
                    if should_attach_file:
                        # Create synthetic document file
                        file_path = self._create_synthetic_document_file(
                            document=document,
                            deal=deal,
                            file_type=random.choice(["docx", "pdf"])
                        )
                        if file_path:
                            logger.debug(f"Attached file to document {document.id}: {file_path}")
                    
                    documents.append(document)
                    self._update_status("documents", current=len(documents), progress=len(documents) / (len(deals) * 2))
                    
            except Exception as e:
                error_msg = f"Failed to generate document for deal {deal.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("documents", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("documents", status="completed", completed_at=datetime.utcnow())
        
        return documents
    
    def _create_document_versions(self, documents: List[Document]) -> List[DocumentVersion]:
        """
        Create DocumentVersions (24-48) with version_number, extracted_text, extracted_data (CDM JSON).
        
        Args:
            documents: List of Document objects
            
        Returns:
            List of DocumentVersion objects
        """
        self._update_status("document_versions", status="running", started_at=datetime.utcnow(), total=len(documents) * 1.5, current=0)
        
        document_versions = []
        
        for document in documents:
            try:
                # Check if versions already exist for this document, skip if they do
                existing_versions = self.db.query(DocumentVersion).filter(
                    DocumentVersion.document_id == document.id
                ).all()
                if existing_versions:
                    logger.debug(f"Document versions already exist for document {document.id}, skipping version generation")
                    document_versions.extend(existing_versions)
                    # Ensure current_version_id is set
                    if not document.current_version_id and existing_versions:
                        document.current_version_id = existing_versions[0].id
                    continue
                
                # Generate 1-2 versions per document
                version_count = random.randint(1, 2)
                
                for version_num in range(1, version_count + 1):
                    # Generate CDM data for this document
                    cdm = self._generate_cdm_for_deal(
                        deal_type="loan_application",
                        seed=document.id + version_num
                    )
                    
                    # Generate extracted text (first 5000 chars of simulated document)
                    extracted_text = f"""CREDIT AGREEMENT

This Credit Agreement (the "Agreement") is entered into on {cdm.agreement_date or '2024-01-15'} 
between the parties listed herein.

BORROWER: {document.borrower_name or 'Borrower'}
LEI: {document.borrower_lei or 'N/A'}

FACILITY DETAILS:
"""
                    if cdm.facilities:
                        for facility in cdm.facilities:
                            extracted_text += f"""
Facility Name: {facility.facility_name}
Commitment Amount: {facility.commitment_amount.amount} {facility.commitment_amount.currency}
Maturity Date: {facility.maturity_date}
Interest Rate: {facility.interest_terms.rate_option.benchmark} + {facility.interest_terms.rate_option.spread_bps} bps
"""
                    
                    extracted_text = extracted_text[:5000]  # Limit to 5000 chars
                    
                    # Check if this version already exists, skip if it does
                    existing_version = self.db.query(DocumentVersion).filter(
                        DocumentVersion.document_id == document.id,
                        DocumentVersion.version_number == version_num
                    ).first()
                    if existing_version:
                        logger.debug(f"Document version {version_num} already exists for document {document.id}, skipping")
                        document_versions.append(existing_version)
                        continue
                    
                    # Create document version
                    # Use model_dump_json() and parse to ensure dates are serialized as strings for JSONB storage
                    import json
                    cdm_json = cdm.model_dump_json()
                    cdm_dict = json.loads(cdm_json)
                    # Generate source filename for document version
                    source_filename = f"{document.title.replace(' ', '_')}_v{version_num}.pdf"
                    if document.deal:
                        # Try to get actual file path from deal folder
                        try:
                            deal_files = self.file_storage.get_deal_documents(
                                user_id=document.deal.applicant_id,
                                deal_id=document.deal.deal_id,
                                subdirectory="documents"
                            )
                            # Find file matching this document
                            for file_info in deal_files:
                                if str(document.id) in file_info.get("filename", ""):
                                    source_filename = file_info.get("filename", source_filename)
                                    break
                        except Exception as e:
                            logger.debug(f"Could not retrieve file info for document {document.id}: {e}")
                    
                    doc_version = DocumentVersion(
                        document_id=document.id,
                        version_number=version_num,
                        original_text=extracted_text,  # Use original_text, not extracted_text
                        extracted_data=cdm_dict,
                        source_filename=source_filename,  # Populate source_filename
                        extraction_method=random.choice(["simple", "map_reduce", "structured"]),
                        created_by=document.uploaded_by,  # Set created_by from document
                        created_at=document.created_at + timedelta(days=random.randint(0, 2))
                    )
                    
                    self.db.add(doc_version)
                    # Flush to get the ID before updating document's current_version_id
                    self.db.flush()
                    document_versions.append(doc_version)
                    
                    # Update document's current_version_id if this is version 1
                    if version_num == 1:
                        document.current_version_id = doc_version.id
                        # Also set source_cdm_data for easier access
                        document.source_cdm_data = cdm_dict
                    
                    self._update_status("document_versions", current=len(document_versions), progress=len(document_versions) / (len(documents) * 1.5))
                
            except Exception as e:
                error_msg = f"Failed to create document version for document {document.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("document_versions", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("document_versions", status="completed", completed_at=datetime.utcnow())
        
        return document_versions
    
    def create_document_revision_history(self, document_ids: List[int]) -> List[DocumentVersion]:
        """
        Generate document revision history (30% of documents get version 2, 3) with modified CDM data.
        
        Args:
            document_ids: List of document IDs to create revision history for
            
        Returns:
            List of additional DocumentVersion objects (versions 2, 3)
        """
        self._update_status("document_revisions", status="running", started_at=datetime.utcnow(), total=len(document_ids), current=0)
        
        additional_versions = []
        
        # Handle empty document list
        if not document_ids:
            self._update_status("document_revisions", status="skipped", completed_at=datetime.utcnow())
            return additional_versions
        
        # Select 30% of documents for revision history
        sample_size = max(1, int(len(document_ids) * 0.30))
        # Ensure sample size doesn't exceed population
        sample_size = min(sample_size, len(document_ids))
        documents_to_revise = random.sample(document_ids, sample_size)
        
        for document_id in documents_to_revise:
            try:
                document = self.db.query(Document).filter(Document.id == document_id).first()
                if not document:
                    continue
                
                # Get existing versions to find max version number
                existing_versions = self.db.query(DocumentVersion).filter(
                    DocumentVersion.document_id == document_id
                ).order_by(DocumentVersion.version_number.desc()).all()
                
                max_version = max([v.version_number for v in existing_versions]) if existing_versions else 1
                
                # Create 1-2 additional versions (version 2, 3)
                num_additional = random.randint(1, 2)
                
                for i in range(num_additional):
                    version_num = max_version + i + 1
                    
                    # Get base CDM from previous version
                    base_version = existing_versions[0] if existing_versions else None
                    base_cdm = None
                    
                    if base_version and base_version.extracted_data:
                        try:
                            from app.models.cdm import CreditAgreement
                            base_cdm = CreditAgreement(**base_version.extracted_data)
                        except Exception:
                            pass
                    
                    if not base_cdm:
                        # Generate new CDM
                        base_cdm = self._generate_cdm_for_deal(deal_type="loan_application", seed=document_id + version_num)
                    else:
                        # Modify CDM slightly (simulate edits)
                        # Change amounts, dates, or other fields
                        if base_cdm.facilities and len(base_cdm.facilities) > 0:
                            # Modify commitment amount slightly
                            original_amount = base_cdm.facilities[0].commitment_amount.amount
                            modification_factor = random.uniform(0.95, 1.05)  # ±5% change
                            new_amount = Decimal(str(float(original_amount) * modification_factor))
                            base_cdm.facilities[0].commitment_amount.amount = new_amount
                        
                        # Update agreement date
                        if base_cdm.agreement_date:
                            base_cdm.agreement_date = base_cdm.agreement_date + timedelta(days=random.randint(-5, 5))
                    
                    # Generate modified extracted text
                    extracted_text = f"""CREDIT AGREEMENT (REVISED VERSION {version_num})

This Credit Agreement (the "Agreement") is entered into on {base_cdm.agreement_date or '2024-01-15'} 
between the parties listed herein.

BORROWER: {document.borrower_name or 'Borrower'}
LEI: {document.borrower_lei or 'N/A'}

FACILITY DETAILS (REVISED):
"""
                    if base_cdm.facilities:
                        for facility in base_cdm.facilities:
                            extracted_text += f"""
Facility Name: {facility.facility_name}
Commitment Amount: {facility.commitment_amount.amount} {facility.commitment_amount.currency}
Maturity Date: {facility.maturity_date}
Interest Rate: {facility.interest_terms.rate_option.benchmark} + {facility.interest_terms.rate_option.spread_bps} bps
"""
                    
                    extracted_text = extracted_text[:5000]
                    
                    # Check if this version already exists, skip if it does
                    existing_version = self.db.query(DocumentVersion).filter(
                        DocumentVersion.document_id == document_id,
                        DocumentVersion.version_number == version_num
                    ).first()
                    if existing_version:
                        logger.debug(f"Document version {version_num} already exists for document {document_id}, skipping")
                        additional_versions.append(existing_version)
                        continue
                    
                    # Create document version
                    # Use model_dump_json() and parse to ensure dates are serialized as strings for JSONB storage
                    import json
                    cdm_json = base_cdm.model_dump_json()
                    cdm_dict = json.loads(cdm_json)
                    doc_version = DocumentVersion(
                        document_id=document_id,
                        version_number=version_num,
                        original_text=extracted_text,  # Use original_text, not extracted_text
                        extracted_data=cdm_dict,
                        created_at=document.created_at + timedelta(days=random.randint(1, 5))
                    )
                    
                    self.db.add(doc_version)
                    additional_versions.append(doc_version)
                    
                    # Link to workflow state transitions
                    # Version 2 → under_review state (if exists)
                    # Version 3 → approved state (if exists)
                    workflow = self.db.query(Workflow).filter(Workflow.document_id == document_id).first()
                    if workflow:
                        if version_num == 2 and workflow.state == WorkflowState.UNDER_REVIEW.value:
                            workflow.submitted_at = doc_version.created_at
                        elif version_num == 3 and workflow.state == WorkflowState.APPROVED.value:
                            workflow.approved_at = doc_version.created_at
                    
                    self._update_status("document_revisions", current=len(additional_versions), progress=len(additional_versions) / len(documents_to_revise))
                
            except Exception as e:
                error_msg = f"Failed to create revision history for document {document_id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("document_revisions", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("document_revisions", status="completed", completed_at=datetime.utcnow())
        
        return additional_versions
    
    def _create_workflows_for_documents(self, documents: List[Document]) -> List[Workflow]:
        """
        Create Workflows (18-36) with state distribution, assigned_to, approved_by, timestamps.
        
        Args:
            documents: List of Document objects
            
        Returns:
            List of Workflow objects
        """
        self._update_status("workflows", status="running", started_at=datetime.utcnow(), total=len(documents), current=0)
        
        workflows = []
        # Get reviewers and approvers
        reviewers = self.db.query(User).filter(
            User.role.in_([UserRole.LAW_OFFICER.value, UserRole.REVIEWER.value, UserRole.ADMIN.value])
        ).all()
        creators = self.db.query(User).filter(
            User.role.in_([UserRole.BANKER.value, UserRole.ANALYST.value])
        ).all()
        
        if not reviewers:
            reviewers = self.db.query(User).limit(3).all()
        if not creators:
            creators = self.db.query(User).limit(2).all()
        
        state_distribution = {
            WorkflowState.DRAFT.value: 0.30,
            WorkflowState.UNDER_REVIEW.value: 0.40,
            WorkflowState.APPROVED.value: 0.20,
            WorkflowState.PUBLISHED.value: 0.10
        }
        
        for document in documents:
            try:
                # Check if workflow already exists for this document, skip if it does
                existing_workflow = self.db.query(Workflow).filter(Workflow.document_id == document.id).first()
                if existing_workflow:
                    logger.debug(f"Workflow already exists for document {document.id}, skipping workflow creation")
                    workflows.append(existing_workflow)
                    continue
                
                # Select state based on distribution
                state = random.choices(
                    list(state_distribution.keys()),
                    weights=list(state_distribution.values())
                )[0]
                
                # Assign based on state
                assigned_to = None
                if state == WorkflowState.UNDER_REVIEW.value:
                    assigned_to = random.choice(reviewers).id if reviewers else None
                elif state == WorkflowState.DRAFT.value:
                    assigned_to = random.choice(creators).id if creators else None
                elif state in [WorkflowState.APPROVED.value, WorkflowState.PUBLISHED.value]:
                    assigned_to = random.choice(reviewers).id if reviewers else None
                
                # Generate timestamps
                submitted_at = None
                approved_at = None
                published_at = None
                
                if state in [WorkflowState.UNDER_REVIEW.value, WorkflowState.APPROVED.value, WorkflowState.PUBLISHED.value]:
                    submitted_at = document.created_at + timedelta(days=random.randint(1, 3))
                
                if state in [WorkflowState.APPROVED.value, WorkflowState.PUBLISHED.value]:
                    approved_at = submitted_at + timedelta(days=random.randint(1, 7)) if submitted_at else None
                
                if state == WorkflowState.PUBLISHED.value:
                    published_at = approved_at + timedelta(days=random.randint(1, 3)) if approved_at else None
                
                # Select approver
                approved_by = random.choice(reviewers).id if reviewers and state in [WorkflowState.APPROVED.value, WorkflowState.PUBLISHED.value] else None
                
                # Priority
                priority = random.choices(
                    ["normal", "high", "urgent"],
                    weights=[0.70, 0.25, 0.05]
                )[0]
                
                # Due date
                due_date = None
                if state == WorkflowState.UNDER_REVIEW.value and submitted_at:
                    due_date = submitted_at + timedelta(days=random.randint(7, 14))
                
                # Create workflow
                workflow = Workflow(
                    document_id=document.id,
                    state=state,
                    assigned_to=assigned_to,
                    submitted_at=submitted_at,
                    approved_at=approved_at,
                    approved_by=approved_by,
                    published_at=published_at,
                    priority=priority,
                    due_date=due_date
                )
                
                self.db.add(workflow)
                workflows.append(workflow)
                self._update_status("workflows", current=len(workflows), progress=len(workflows) / len(documents))
                
            except Exception as e:
                error_msg = f"Failed to create workflow for document {document.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("workflows", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("workflows", status="completed", completed_at=datetime.utcnow())
        
        return workflows
    
    def _generate_loan_assets_for_deals(self, deals: List[Deal]) -> List[LoanAsset]:
        """
        Generate LoanAssets (4-6) for sustainability-linked deals.
        
        Args:
            deals: List of Deal objects
            
        Returns:
            List of LoanAsset objects
        """
        # Filter sustainability-linked deals
        sustainability_deals = [
            deal for deal in deals
            if deal.deal_data and deal.deal_data.get("sustainability_linked", False)
        ]
        
        if not sustainability_deals:
            return []
        
        self._update_status("loan_assets", status="running", started_at=datetime.utcnow(), total=len(sustainability_deals), current=0)
        
        loan_assets = []
        
        for deal in sustainability_deals:
            try:
                # Get ESG targets from deal
                esg_targets = deal.deal_data.get("esg_targets", {}) if deal.deal_data else {}
                ndvi_threshold = esg_targets.get("ndvi_threshold", 0.75)
                
                # Generate realistic address
                addresses = [
                    "123 Industrial Way, Detroit, MI 48201",
                    "456 Farm Road, Napa, CA 94558",
                    "789 Agricultural Blvd, Fresno, CA 93721",
                    "321 Green Valley Lane, Austin, TX 78701"
                ]
                collateral_address = random.choice(addresses)
                
                # Generate coordinates (rough geocoding)
                geo_coords = {
                    "123 Industrial Way, Detroit, MI 48201": (42.3314, -83.0458),
                    "456 Farm Road, Napa, CA 94558": (38.2975, -122.2869),
                    "789 Agricultural Blvd, Fresno, CA 93721": (36.7378, -119.7871),
                    "321 Green Valley Lane, Austin, TX 78701": (30.2672, -97.7431)
                }
                geo_lat, geo_lon = geo_coords.get(collateral_address, (40.7128, -74.0060))
                
                # Generate NDVI score
                last_verified_score = random.uniform(0.50, 0.95)
                
                # Calculate risk status
                if last_verified_score >= ndvi_threshold:
                    risk_status = "COMPLIANT"
                elif last_verified_score >= ndvi_threshold * 0.9:
                    risk_status = "WARNING"
                else:
                    risk_status = "BREACH"
                
                # Get interest rate
                base_interest_rate = deal.deal_data.get("interest_rate", 5.25) if deal.deal_data else 5.25
                penalty_bps = esg_targets.get("penalty_bps", 25) if isinstance(esg_targets, dict) else 25
                current_interest_rate = base_interest_rate + (penalty_bps / 10000 if risk_status == "BREACH" else 0)
                
                # Generate legal reality data
                # Original text from document extraction (simulated)
                loan_amount = deal.deal_data.get('loan_amount', 1000000) if deal.deal_data else 1000000
                term_years = deal.deal_data.get('term_years', 5) if deal.deal_data else 5
                original_text = f"""Credit Agreement for {deal.deal_id}
                
This sustainability-linked credit agreement establishes a loan facility with the following terms:
- Loan Amount: ${loan_amount:,.0f}
- Interest Rate: {base_interest_rate}% base rate
- Term: {term_years} years
- Collateral: {collateral_address}
- ESG Target: NDVI threshold of {ndvi_threshold:.2f}
- Penalty: {penalty_bps} basis points for non-compliance

The borrower agrees to maintain the specified NDVI threshold and will be subject to quarterly verification."""
                
                # Generate mock legal vector (1536-dim embeddings, similar to OpenAI text-embedding-3-large)
                legal_vector = [random.uniform(-1.0, 1.0) for _ in range(1536)]
                
                # Generate physical reality data
                # Mock satellite snapshot URL
                satellite_snapshot_url = f"https://josephrp.github.io/creditnexus/snapshots/{deal.deal_id}/{datetime.utcnow().strftime('%Y%m%d')}.tif"
                
                # Generate mock geo vector (512-dim image embeddings)
                geo_vector = [random.uniform(-1.0, 1.0) for _ in range(512)]
                
                # Enhanced SPT data with CDM compliance and measurement history
                spt_data = {
                    "target_type": "NDVI",
                    "threshold": ndvi_threshold,
                    "measurement_frequency": "Quarterly",
                    "penalty_bps": penalty_bps,
                    "cdm_compliant": True,
                    "measurement_history": [
                        {
                            "date": (datetime.utcnow() - timedelta(days=90*i)).isoformat(),
                            "value": max(0.0, min(1.0, last_verified_score + random.uniform(-0.1, 0.1))),
                            "verified_by": "Sentinel-2B",
                            "confidence": random.uniform(0.85, 0.98)
                        }
                        for i in range(4)  # Last 4 quarters
                    ],
                    "next_verification_date": (datetime.utcnow() + timedelta(days=90)).isoformat()
                }
                
                # Enhanced asset metadata
                verification_error = None
                if random.random() < 0.05:  # 5% of assets have verification errors
                    verification_error = random.choice([
                        "Cloud cover exceeded threshold",
                        "Satellite imagery unavailable",
                        "Geographic coordinates invalid",
                        "Temporary service interruption"
                    ])
                
                asset_metadata = {
                    "verification_method": "Sentinel-2B",
                    "cloud_cover": random.uniform(0.01, 0.10),
                    "classification": random.choice(["AnnualCrop", "Forest", "PermanentCrop"]),
                    "confidence": random.uniform(0.85, 0.98),
                    "penalty_payment_flags": {
                        "has_pending_penalty": risk_status == "BREACH",
                        "penalty_amount": penalty_bps * deal.deal_data.get('loan_amount', 1000000) / 10000 if deal.deal_data and risk_status == "BREACH" else 0,
                        "penalty_applied_date": (datetime.utcnow() - timedelta(days=random.randint(1, 30))).isoformat() if risk_status == "BREACH" else None
                    },
                    "verification_history": [
                        {
                            "date": (datetime.utcnow() - timedelta(days=30*i)).isoformat(),
                            "score": max(0.0, min(1.0, last_verified_score + random.uniform(-0.05, 0.05))),
                            "status": random.choice(["COMPLIANT", "WARNING", "BREACH"]),
                            "method": "Sentinel-2B"
                        }
                        for i in range(3)  # Last 3 verifications
                    ],
                    "satellite_imagery_metadata": {
                        "source": "Sentinel-2B",
                        "resolution": "10m",
                        "acquisition_date": (datetime.utcnow() - timedelta(days=random.randint(1, 30))).isoformat(),
                        "processing_date": datetime.utcnow().isoformat(),
                        "bands": ["B04", "B08", "B11", "B12"],
                        "ndvi_calculation_method": "standard"
                    }
                }
                
                # Generate green finance metrics (synthetic for demo)
                # Location classification based on address pattern
                location_type_map = {
                    "123 Industrial Way, Detroit, MI 48201": "urban",
                    "456 Farm Road, Napa, CA 94558": "rural",
                    "789 Agricultural Blvd, Fresno, CA 93721": "suburban",
                    "321 Green Valley Lane, Austin, TX 78701": "suburban"
                }
                location_type = location_type_map.get(collateral_address, random.choice(["urban", "suburban", "rural"]))
                location_confidence = random.uniform(0.75, 0.95)
                
                # Generate air quality data (synthetic)
                # Urban areas typically have higher AQI, rural lower
                if location_type == "urban":
                    aqi = random.uniform(80, 150)  # Moderate to Unhealthy for Sensitive Groups
                    pm25 = random.uniform(25, 55)
                    pm10 = random.uniform(50, 100)
                    no2 = random.uniform(30, 80)
                elif location_type == "suburban":
                    aqi = random.uniform(50, 100)  # Good to Moderate
                    pm25 = random.uniform(12, 35)
                    pm10 = random.uniform(30, 70)
                    no2 = random.uniform(15, 50)
                else:  # rural
                    aqi = random.uniform(20, 60)  # Good
                    pm25 = random.uniform(5, 20)
                    pm10 = random.uniform(10, 40)
                    no2 = random.uniform(5, 25)
                
                # Generate OSM metrics (synthetic)
                if location_type == "urban":
                    building_count = random.randint(500, 2000)
                    building_density = random.uniform(50, 150)  # buildings/km²
                    road_density = random.uniform(8, 15)  # km/km²
                    green_infrastructure_coverage = random.uniform(0.10, 0.25)  # 10-25%
                elif location_type == "suburban":
                    building_count = random.randint(100, 500)
                    building_density = random.uniform(20, 50)  # buildings/km²
                    road_density = random.uniform(5, 10)  # km/km²
                    green_infrastructure_coverage = random.uniform(0.25, 0.40)  # 25-40%
                else:  # rural
                    building_count = random.randint(10, 100)
                    building_density = random.uniform(2, 20)  # buildings/km²
                    road_density = random.uniform(2, 6)  # km/km²
                    green_infrastructure_coverage = random.uniform(0.40, 0.70)  # 40-70%
                
                # Calculate sustainability components
                # Normalize NDVI to 0-1 (already is)
                vegetation_health = last_verified_score
                
                # Normalize AQI to 0-1 (inverse: lower AQI = better = higher score)
                # AQI 0-50 = 1.0, 51-100 = 0.8, 101-150 = 0.6, 151-200 = 0.4, 201+ = 0.2
                if aqi <= 50:
                    air_quality_score = 1.0
                elif aqi <= 100:
                    air_quality_score = 0.8
                elif aqi <= 150:
                    air_quality_score = 0.6
                elif aqi <= 200:
                    air_quality_score = 0.4
                else:
                    air_quality_score = 0.2
                
                # Urban activity (inverse: less activity = better for sustainability)
                # Based on building/road density
                activity_score = max(0.0, 1.0 - (building_density / 200.0) - (road_density / 20.0))
                
                # Green infrastructure (direct: more green = better)
                green_infra_score = green_infrastructure_coverage
                
                # Pollution levels (inverse: less pollution = better)
                pollution_score = max(0.0, 1.0 - (aqi / 300.0))
                
                # Calculate composite sustainability score (weighted average)
                # Using default weights from config
                sustainability_components = {
                    "vegetation_health": vegetation_health,
                    "air_quality": air_quality_score,
                    "urban_activity": activity_score,
                    "green_infrastructure": green_infra_score,
                    "pollution_levels": pollution_score
                }
                
                # Default weights (matching config defaults)
                weights = {
                    "vegetation_health": 0.30,
                    "air_quality": 0.25,
                    "urban_activity": 0.15,
                    "green_infrastructure": 0.20,
                    "pollution_levels": 0.10
                }
                
                composite_sustainability_score = (
                    sustainability_components["vegetation_health"] * weights["vegetation_health"] +
                    sustainability_components["air_quality"] * weights["air_quality"] +
                    sustainability_components["urban_activity"] * weights["urban_activity"] +
                    sustainability_components["green_infrastructure"] * weights["green_infrastructure"] +
                    sustainability_components["pollution_levels"] * weights["pollution_levels"]
                )
                
                # Build green finance metrics dict
                green_finance_metrics = {
                    "location_type": location_type,
                    "location_confidence": location_confidence,
                    "osm_metrics": {
                        "building_count": building_count,
                        "building_density": building_density,
                        "road_density": road_density,
                        "green_infrastructure_coverage": green_infrastructure_coverage
                    },
                    "air_quality": {
                        "pm25": pm25,
                        "pm10": pm10,
                        "no2": no2,
                        "data_source": "synthetic_demo"
                    },
                    "sustainability_components": sustainability_components
                }
                
                # Set verification timestamp
                last_verified_at = datetime.utcnow() - timedelta(days=random.randint(1, 30))
                
                # Create loan asset with all enhanced fields
                loan_asset = LoanAsset(
                    loan_id=deal.deal_id,
                    # Legal Reality
                    original_text=original_text,
                    legal_vector=legal_vector,
                    # Physical Reality
                    satellite_snapshot_url=satellite_snapshot_url,
                    geo_vector=geo_vector,
                    # Basic Info
                    collateral_address=collateral_address,
                    geo_lat=geo_lat,
                    geo_lon=geo_lon,
                    # SPT Data (enhanced)
                    spt_data=spt_data,
                    spt_threshold=ndvi_threshold,
                    # Verification
                    last_verified_score=last_verified_score,
                    last_verified_at=last_verified_at,
                    verification_error=verification_error,
                    risk_status=risk_status,
                    # Interest Rates
                    base_interest_rate=float(base_interest_rate),
                    current_interest_rate=float(current_interest_rate),
                    penalty_bps=penalty_bps if risk_status == "BREACH" else 0,
                    # Metadata (enhanced)
                    asset_metadata=asset_metadata,
                    # Green Finance Metrics (always set)
                    location_type=location_type,
                    air_quality_index=aqi,
                    composite_sustainability_score=composite_sustainability_score,
                    green_finance_metrics=green_finance_metrics
                )
                
                # Add to session - SQLModel works with SQLAlchemy sessions
                try:
                    self.db.add(loan_asset)
                    self.db.flush()  # Flush to get ID and validate
                    loan_assets.append(loan_asset)
                    self._update_status("loan_assets", current=len(loan_assets), progress=len(loan_assets) / len(sustainability_deals))
                except Exception as db_error:
                    # Rollback the failed add
                    self.db.rollback()
                    error_msg = f"Failed to add loan asset to database for deal {deal.id}: {str(db_error)}"
                    logger.error(error_msg, exc_info=True)
                    logger.error(f"LoanAsset data: loan_id={deal.deal_id}, risk_status={risk_status}, base_rate={base_interest_rate}, current_rate={current_interest_rate}")
                    self._update_status("loan_assets", errors=[error_msg])
                    continue
                
            except Exception as e:
                error_msg = f"Failed to generate loan asset for deal {deal.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                logger.error(f"Exception type: {type(e).__name__}, args: {e.args}")
                self._update_status("loan_assets", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("loan_assets", status="completed", completed_at=datetime.utcnow())
        
        return loan_assets
    
    def create_demo_securitization_pools(
        self,
        deals: List[Deal],
        loan_assets: List[LoanAsset],
        num_pools: int = 3
    ) -> List[SecuritizationPool]:
        """
        Create demo securitization pools with tranches and assets.
        
        Args:
            deals: List of Deal objects to pool
            loan_assets: List of LoanAsset objects to pool
            num_pools: Number of pools to create (default: 3)
            
        Returns:
            List of SecuritizationPool objects
        """
        if not deals and not loan_assets:
            logger.warning("No deals or loan assets available for securitization")
            return []
        
        self._update_status("securitization_pools", status="running", started_at=datetime.utcnow(), total=num_pools, current=0)
        
        # Get eligible users for originator and trustee roles
        originators = self.db.query(User).filter(User.role == UserRole.BANKER.value).all()
        trustees = self.db.query(User).filter(User.role.in_([UserRole.BANKER.value, UserRole.LAW_OFFICER.value])).all()
        
        if not originators:
            originators = self.db.query(User).limit(2).all()
        if not trustees:
            trustees = self.db.query(User).limit(2).all()
        
        pools = []
        pool_counter = 1
        now = datetime.utcnow()
        
        # Select eligible assets (active deals and loan assets)
        eligible_deals = [d for d in deals if d.status in [DealStatus.ACTIVE.value, DealStatus.APPROVED.value]]
        eligible_loan_assets = loan_assets  # All loan assets are eligible
        
        if not eligible_deals and not eligible_loan_assets:
            logger.warning("No eligible deals or loan assets for securitization")
            return []
        
        for i in range(num_pools):
            try:
                # Generate pool_id (format: POOL-YYYY-MM-XXX)
                pool_id = f"POOL-{now.year}-{now.month:02d}-{pool_counter:03d}"
                
                # Check if pool already exists
                existing_pool = self.db.query(SecuritizationPool).filter(SecuritizationPool.pool_id == pool_id).first()
                if existing_pool:
                    logger.debug(f"Pool {pool_id} already exists, skipping")
                    pool_counter += 1
                    continue
                
                # Select originator and trustee
                originator = random.choice(originators) if originators else None
                trustee = random.choice(trustees) if trustees else None
                
                # Select assets for this pool (3-8 assets per pool)
                num_assets = random.randint(3, min(8, len(eligible_deals) + len(eligible_loan_assets)))
                selected_deals = random.sample(eligible_deals, min(num_assets // 2, len(eligible_deals))) if eligible_deals else []
                remaining_slots = num_assets - len(selected_deals)
                selected_loan_assets = random.sample(eligible_loan_assets, min(remaining_slots, len(eligible_loan_assets))) if eligible_loan_assets else []
                
                # Calculate total pool value
                deal_values = [Decimal(str(d.deal_data.get("loan_amount", 1000000))) if d.deal_data else Decimal("1000000") for d in selected_deals]
                loan_asset_values = [Decimal(str(la.asset_metadata.get("estimated_value", 500000))) if la.asset_metadata and "estimated_value" in la.asset_metadata else Decimal("500000") for la in selected_loan_assets]
                total_pool_value = sum(deal_values) + sum(loan_asset_values)
                
                # Generate pool name
                pool_types = ["ABS", "CLO", "MBS"]
                pool_type = random.choice(pool_types)
                pool_name = f"Q{((now.month - 1) // 3) + 1} {now.year} {pool_type} Pool {pool_counter}"
                
                # Generate CDM-compliant payload
                cdm_payload = {
                    "pool_id": pool_id,
                    "pool_name": pool_name,
                    "pool_type": pool_type,
                    "originator": {
                        "id": f"user_{originator.id}" if originator else "demo_originator",
                        "name": originator.display_name if originator else "Demo Originator",
                        "role": "Originator",
                        "lei": originator.profile_data.get("company", {}).get("lei") if originator and originator.profile_data else None
                    },
                    "trustee": {
                        "id": f"user_{trustee.id}" if trustee else "demo_trustee",
                        "name": trustee.display_name if trustee else "Demo Trustee",
                        "role": "Trustee",
                        "lei": trustee.profile_data.get("company", {}).get("lei") if trustee and trustee.profile_data else None
                    },
                    "total_pool_value": {
                        "amount": float(total_pool_value),
                        "currency": "USD"
                    },
                    "creation_date": now.date().isoformat(),
                    "effective_date": (now + timedelta(days=30)).date().isoformat(),
                    "maturity_date": (now + timedelta(days=365*5)).date().isoformat()
                }
                
                # Set status distribution
                rand = random.random()
                if rand < 0.20:
                    status = "draft"
                elif rand < 0.50:
                    status = "pending_notarization"
                elif rand < 0.80:
                    status = "notarized"
                elif rand < 0.95:
                    status = "filed"
                else:
                    status = "active"
                
                # Create pool
                pool = SecuritizationPool(
                    pool_id=pool_id,
                    pool_name=pool_name,
                    pool_type=pool_type,
                    originator_id=originator.id if originator else None,
                    trustee_id=trustee.id if trustee else None,
                    total_pool_value=total_pool_value,
                    currency="USD",
                    cdm_payload=cdm_payload,
                    cdm_data={
                        "payment_schedule": "Monthly",
                        "interest_calculation_method": "Actual/360"
                    },
                    status=status,
                    notarized_at=now if status in ["notarized", "filed", "active"] else None,
                    filed_at=now if status in ["filed", "active"] else None
                )
                
                self.db.add(pool)
                self.db.flush()  # Get pool.id
                
                # Generate tranches
                tranches = self._generate_securitization_tranches(pool, total_pool_value)
                
                # Generate pool assets
                pool_assets = self._generate_pool_assets(pool, selected_deals, selected_loan_assets, total_pool_value)
                
                # Generate regulatory filings (for notarized/filed pools)
                if status in ["notarized", "filed", "active"]:
                    filings = self._generate_regulatory_filings(pool)
                
                # Generate notarization records (for notarized pools)
                if status in ["notarized", "filed", "active"]:
                    notarizations = self._generate_notarization_records(pool)
                
                pools.append(pool)
                pool_counter += 1
                
                self._update_status("securitization_pools", current=len(pools), progress=len(pools) / num_pools)
                
            except Exception as e:
                error_msg = f"Failed to create securitization pool {pool_counter}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("securitization_pools", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("securitization_pools", status="completed", completed_at=datetime.utcnow())
        
        return pools
    
    def _generate_securitization_tranches(
        self,
        pool: SecuritizationPool,
        total_pool_value: Decimal
    ) -> List[SecuritizationTranche]:
        """
        Generate tranches for a securitization pool.
        
        Args:
            pool: SecuritizationPool object
            total_pool_value: Total value of the pool
            
        Returns:
            List of SecuritizationTranche objects
        """
        tranches = []
        
        # Tranche structure: Senior (60-70%), Mezzanine (20-30%), Equity (5-15%)
        senior_pct = random.uniform(0.60, 0.70)
        mezzanine_pct = random.uniform(0.20, 0.30)
        equity_pct = 1.0 - senior_pct - mezzanine_pct
        
        # Ensure equity is at least 5%
        if equity_pct < 0.05:
            equity_pct = 0.05
            senior_pct = 0.65
            mezzanine_pct = 0.30
        
        tranche_configs = [
            {
                "class": "Senior",
                "percentage": senior_pct,
                "risk_ratings": ["AAA", "AA"],
                "interest_rate_range": (0.03, 0.05),
                "priority": 1
            },
            {
                "class": "Mezzanine",
                "percentage": mezzanine_pct,
                "risk_ratings": ["A", "BBB"],
                "interest_rate_range": (0.06, 0.09),
                "priority": 2
            },
            {
                "class": "Equity",
                "percentage": equity_pct,
                "risk_ratings": ["BB", "B"],
                "interest_rate_range": (0.10, 0.15),
                "priority": 3
            }
        ]
        
        for idx, config in enumerate(tranche_configs):
            tranche_size = total_pool_value * Decimal(str(config["percentage"]))
            risk_rating = random.choice(config["risk_ratings"])
            interest_rate = Decimal(str(random.uniform(*config["interest_rate_range"])))
            
            tranche_id = f"{pool.pool_id}-{config['class']}"
            tranche_name = f"{config['class']} Tranche"
            
            # Generate CDM tranche data
            cdm_data = {
                "tranche_id": tranche_id,
                "tranche_name": tranche_name,
                "tranche_class": config["class"],
                "size": {
                    "amount": float(tranche_size),
                    "currency": "USD"
                },
                "interest_rate": float(interest_rate),
                "risk_rating": risk_rating,
                "payment_priority": config["priority"]
            }
            
            tranche = SecuritizationTranche(
                pool_id=pool.id,
                tranche_id=tranche_id,
                tranche_name=tranche_name,
                tranche_class=config["class"],
                size=tranche_size,
                currency="USD",
                interest_rate=interest_rate,
                risk_rating=risk_rating,
                payment_priority=config["priority"],
                principal_remaining=tranche_size,
                interest_accrued=Decimal("0"),
                token_id=f"TOKEN-{pool.pool_id}-{config['class']}-{idx+1}",
                owner_wallet_address=f"0x{''.join([random.choice('0123456789abcdef') for _ in range(40)])}",
                cdm_data=cdm_data
            )
            
            self.db.add(tranche)
            tranches.append(tranche)
        
        return tranches
    
    def _generate_pool_assets(
        self,
        pool: SecuritizationPool,
        deals: List[Deal],
        loan_assets: List[LoanAsset],
        total_pool_value: Decimal
    ) -> List[SecuritizationPoolAsset]:
        """
        Generate pool assets linking deals and loan assets to the pool.
        
        Args:
            pool: SecuritizationPool object
            deals: List of Deal objects
            loan_assets: List of LoanAsset objects
            total_pool_value: Total value of the pool
            
        Returns:
            List of SecuritizationPoolAsset objects
        """
        pool_assets = []
        
        # Add deals
        for deal in deals:
            asset_value = Decimal(str(deal.deal_data.get("loan_amount", 1000000))) if deal.deal_data else Decimal("1000000")
            allocation_percentage = (asset_value / total_pool_value) * 100 if total_pool_value > 0 else 0
            
            pool_asset = SecuritizationPoolAsset(
                pool_id=pool.id,
                deal_id=deal.id,
                loan_asset_id=None,
                asset_type="deal",
                asset_id=deal.deal_id,
                asset_value=asset_value,
                currency="USD",
                allocation_percentage=Decimal(str(allocation_percentage)),
                allocation_amount=asset_value
            )
            
            self.db.add(pool_asset)
            pool_assets.append(pool_asset)
        
        # Add loan assets
        for loan_asset in loan_assets:
            # Try to get value from metadata, otherwise estimate
            if loan_asset.asset_metadata and "estimated_value" in loan_asset.asset_metadata:
                asset_value = Decimal(str(loan_asset.asset_metadata["estimated_value"]))
            else:
                asset_value = Decimal("500000")  # Default estimate
            
            allocation_percentage = (asset_value / total_pool_value) * 100 if total_pool_value > 0 else 0
            
            pool_asset = SecuritizationPoolAsset(
                pool_id=pool.id,
                deal_id=None,
                loan_asset_id=loan_asset.id,
                asset_type="loan_asset",
                asset_id=loan_asset.loan_id,
                asset_value=asset_value,
                currency="USD",
                allocation_percentage=Decimal(str(allocation_percentage)),
                allocation_amount=asset_value
            )
            
            self.db.add(pool_asset)
            pool_assets.append(pool_asset)
        
        return pool_assets
    
    def _generate_regulatory_filings(
        self,
        pool: SecuritizationPool
    ) -> List[RegulatoryFiling]:
        """
        Generate regulatory filings for a securitization pool.
        
        Args:
            pool: SecuritizationPool object
            
        Returns:
            List of RegulatoryFiling objects
        """
        filings = []
        
        filing_types = [
            {"type": "SEC_10D", "body": "SEC", "status": "accepted"},
            {"type": "PROSPECTUS", "body": "SEC", "status": "accepted"},
            {"type": "PSA", "body": "SEC", "status": "accepted"},
            {"type": "TRUST_AGREEMENT", "body": "SEC", "status": "accepted"}
        ]
        
        # Generate 2-3 filings per pool
        selected_filings = random.sample(filing_types, min(random.randint(2, 3), len(filing_types)))
        
        for filing_config in selected_filings:
            filing_number = f"{filing_config['type']}-{pool.pool_id}-{random.randint(1000, 9999)}"
            
            filing = RegulatoryFiling(
                pool_id=pool.id,
                filing_type=filing_config["type"],
                regulatory_body=filing_config["body"],
                filing_number=filing_number,
                status=filing_config["status"],
                document_path=f"storage/securitization/{pool.pool_id}/filings/{filing_config['type']}.pdf",
                filed_at=pool.filed_at or pool.notarized_at or datetime.utcnow(),
                accepted_at=pool.filed_at or pool.notarized_at or datetime.utcnow(),
                filing_metadata={
                    "receipt_number": filing_number,
                    "filing_date": (pool.filed_at or pool.notarized_at or datetime.utcnow()).isoformat()
                }
            )
            
            self.db.add(filing)
            filings.append(filing)
        
        return filings
    
    def _generate_notarization_records(
        self,
        pool: SecuritizationPool
    ) -> List[NotarizationRecord]:
        """
        Generate notarization records for a securitization pool.
        
        Args:
            pool: SecuritizationPool object
            
        Returns:
            List of NotarizationRecord objects
        """
        import hashlib
        import json
        
        # Generate notarization hash from CDM payload
        cdm_json = json.dumps(pool.cdm_payload, sort_keys=True)
        notarization_hash = hashlib.sha256(cdm_json.encode()).hexdigest()
        
        # Generate required signers (originator and trustee wallets)
        required_signers = []
        if pool.originator:
            required_signers.append(f"0x{''.join([random.choice('0123456789abcdef') for _ in range(40)])}")
        if pool.trustee:
            required_signers.append(f"0x{''.join([random.choice('0123456789abcdef') for _ in range(40)])}")
        
        # Generate signatures (all signers have signed for notarized pools)
        signatures = []
        for signer in required_signers:
            signatures.append({
                "wallet_address": signer,
                "signature": f"0x{''.join([random.choice('0123456789abcdef') for _ in range(128)])}",
                "signed_at": (pool.notarized_at or datetime.utcnow()).isoformat()
            })
        
        notarization = NotarizationRecord(
            deal_id=None,  # Securitization pools don't have deal_id
            notarization_hash=notarization_hash,
            required_signers=required_signers,
            signatures=signatures,
            status="completed",
            completed_at=pool.notarized_at or datetime.utcnow(),
            cdm_event_id=f"CDM-EVENT-{pool.pool_id}"
        )
        
        self.db.add(notarization)
        
        return [notarization]
    
    def generate_documents_from_templates(self, deal_id: int) -> List[GeneratedDocument]:
        """
        Generate documents from templates for approved deals (12-24 GeneratedDocuments).
        
        Args:
            deal_id: Deal ID to generate documents for
            
        Returns:
            List of GeneratedDocument objects
        """
        deal = self.db.query(Deal).filter(Deal.id == deal_id).first()
        if not deal:
            raise ValueError(f"Deal {deal_id} not found")
        
        # Only generate for approved deals
        if deal.status not in [DealStatus.APPROVED.value, DealStatus.ACTIVE.value]:
            logger.info(f"Skipping template generation for deal {deal_id} (status: {deal.status})")
            return []
        
        self._update_status("generated_documents", status="running", started_at=datetime.utcnow(), total=2, current=0)
        
        # Get available templates
        from app.db.models import LMATemplate
        templates = self.db.query(LMATemplate).filter(LMATemplate.is_active == True).all()
        
        if not templates:
            logger.warning("No active templates found for document generation")
            return []
        
        # Select 1-2 templates per deal
        selected_templates = random.sample(templates, min(random.randint(1, 2), len(templates)))
        
        # Get primary document for CDM data
        primary_doc = self.db.query(Document).filter(
            Document.deal_id == deal_id
        ).order_by(Document.created_at.asc()).first()
        
        if not primary_doc:
            logger.warning(f"No primary document found for deal {deal_id}")
            return []
        
        # Get CDM data from primary document's version
        cdm_data = None
        if primary_doc.current_version_id:
            doc_version = self.db.query(DocumentVersion).filter(
                DocumentVersion.id == primary_doc.current_version_id
            ).first()
            if doc_version and doc_version.extracted_data:
                try:
                    from app.models.cdm import CreditAgreement
                    cdm_data = CreditAgreement(**doc_version.extracted_data)
                except Exception as e:
                    logger.warning(f"Failed to parse CDM data from document version: {e}")
        
        if not cdm_data:
            # Generate CDM if not available
            cdm_data = self._generate_cdm_for_deal(deal_type=deal.deal_type, seed=deal.id)
        
        # Get demo users for document creators
        creators = self.db.query(User).filter(
            User.role.in_([UserRole.BANKER.value, UserRole.LAW_OFFICER.value])
        ).all()
        if not creators:
            creators = self.db.query(User).limit(2).all()
        
        generated_docs = []
        
        try:
            from app.generation.service import DocumentGenerationService
            gen_service = DocumentGenerationService()
            
            for template in selected_templates:
                try:
                    # Generate document
                    generated_doc = gen_service.generate_document(
                        db=self.db,
                        template_id=template.id,
                        cdm_data=cdm_data,
                        user_id=random.choice(creators).id if creators else None,
                        source_document_id=primary_doc.id,
                        deal_id=deal.id
                    )
                    
                    # Update status distribution
                    status_dist = random.choices(
                        [GeneratedDocumentStatus.DRAFT.value, GeneratedDocumentStatus.REVIEW.value, 
                         GeneratedDocumentStatus.APPROVED.value, GeneratedDocumentStatus.EXECUTED.value],
                        weights=[0.40, 0.40, 0.15, 0.05]
                    )[0]
                    generated_doc.status = status_dist
                    
                    # Update file path to demo folder
                    if generated_doc.file_path:
                        # Ensure it's in the demo deal folder
                        if "storage/deals/demo" not in generated_doc.file_path:
                            generated_doc.file_path = f"storage/deals/demo/{deal.deal_id}/generated/{template.template_code}-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
                    
                    # Update generation summary
                    if generated_doc.generation_summary:
                        generated_doc.generation_summary["cache_hits"] = random.randint(0, 3)
                        generated_doc.generation_summary["generation_time_seconds"] = random.uniform(2.5, 8.0)
                    
                    generated_docs.append(generated_doc)
                    self._update_status("generated_documents", current=len(generated_docs), progress=len(generated_docs) / len(selected_templates))
                    
                except Exception as e:
                    error_msg = f"Failed to generate document from template {template.id}: {str(e)}"
                    logger.error(error_msg, exc_info=True)
                    self._update_status("generated_documents", errors=[error_msg])
                    continue
            
            self.db.commit()
            self._update_status("generated_documents", status="completed", completed_at=datetime.utcnow())
            
        except Exception as e:
            error_msg = f"Failed to initialize DocumentGenerationService: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self._update_status("generated_documents", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
        
        return generated_docs
    
    def create_deal_notes(self, deal_ids: List[int]) -> List[DealNote]:
        """
        Create DealNotes (24-60) from different users with note_type distribution.
        
        Args:
            deal_ids: List of deal IDs to create notes for
            
        Returns:
            List of DealNote objects
        """
        self._update_status("deal_notes", status="running", started_at=datetime.utcnow(), total=len(deal_ids) * 3, current=0)
        
        deal_notes = []
        
        # Get users for note creation
        all_users = self.db.query(User).all()
        user_roles = {
            UserRole.BANKER.value: [u for u in all_users if u.role == UserRole.BANKER.value],
            UserRole.LAW_OFFICER.value: [u for u in all_users if u.role == UserRole.LAW_OFFICER.value],
            UserRole.ACCOUNTANT.value: [u for u in all_users if u.role == UserRole.ACCOUNTANT.value],
            UserRole.REVIEWER.value: [u for u in all_users if u.role == UserRole.REVIEWER.value],
            UserRole.APPLICANT.value: [u for u in all_users if u.role == UserRole.APPLICANT.value]
        }
        
        note_templates = {
            "general": [
                "Reviewed credit agreement. Borrower credit score is strong.",
                "Deal progressing well. All documentation in order.",
                "Following up on outstanding items.",
                "Deal status update: All parties aligned."
            ],
            "verification": [
                "Legal review complete. All clauses comply with regulations.",
                "Compliance check passed. No issues identified.",
                "Verification completed successfully.",
                "All required documentation verified."
            ],
            "status_change": [
                "Approved for proceeding to next stage.",
                "Status updated to approved.",
                "Deal moved to active status.",
                "Workflow transition completed."
            ],
            "financial": [
                "Financial analysis shows positive cash flow projections.",
                "Credit analysis complete. Borrower meets all criteria.",
                "Financial review indicates strong credit profile.",
                "Cash flow analysis supports loan approval."
            ]
        }
        
        for deal_id in deal_ids:
            try:
                deal = self.db.query(Deal).filter(Deal.id == deal_id).first()
                if not deal:
                    continue
                
                # Create 2-5 notes per deal
                note_count = random.randint(2, 5)
                
                for i in range(note_count):
                    # Select user role based on distribution
                    role_weights = {
                        UserRole.BANKER.value: 0.30,
                        UserRole.LAW_OFFICER.value: 0.25,
                        UserRole.ACCOUNTANT.value: 0.20,
                        UserRole.REVIEWER.value: 0.15,
                        UserRole.APPLICANT.value: 0.10
                    }
                    
                    selected_role = random.choices(
                        list(role_weights.keys()),
                        weights=list(role_weights.values())
                    )[0]
                    
                    users_for_role = user_roles.get(selected_role, [])
                    if not users_for_role:
                        users_for_role = all_users
                    
                    user = random.choice(users_for_role)
                    
                    # Select note type
                    note_type = random.choices(
                        ["general", "verification", "status_change", "financial"],
                        weights=[0.50, 0.20, 0.20, 0.10]
                    )[0]
                    
                    # Generate note content
                    content = random.choice(note_templates.get(note_type, note_templates["general"]))
                    
                    # Get related document if available
                    related_doc = self.db.query(Document).filter(
                        Document.deal_id == deal_id
                    ).first()
                    
                    note_metadata = {
                        "related_document_id": related_doc.id if related_doc else None,
                        "priority": random.choice(["normal", "high"]),
                        "tags": [note_type, selected_role]
                    }
                    
                    # Generate timestamps across deal lifecycle
                    base_date = deal.created_at
                    note_date = base_date + timedelta(days=random.randint(0, (datetime.utcnow() - base_date).days))
                    
                    # Create note
                    note = DealNote(
                        deal_id=deal_id,
                        user_id=user.id,
                        content=content,
                        note_type=note_type,
                        note_metadata=note_metadata,
                        created_at=note_date,
                        updated_at=note_date + timedelta(days=random.randint(0, 2)) if random.random() < 0.20 else None
                    )
                    
                    self.db.add(note)
                    deal_notes.append(note)
                    self._update_status("deal_notes", current=len(deal_notes), progress=len(deal_notes) / (len(deal_ids) * 3))
                
            except Exception as e:
                error_msg = f"Failed to create notes for deal {deal_id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("deal_notes", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("deal_notes", status="completed", completed_at=datetime.utcnow())
        
        return deal_notes
    
    def create_policy_decisions_for_documents(self, document_ids: List[int]) -> List[PolicyDecision]:
        """
        Create PolicyDecisions (18-36) for documents with decision distribution.
        
        Args:
            document_ids: List of document IDs to create policy decisions for
            
        Returns:
            List of PolicyDecision objects
        """
        self._update_status("policy_decisions", status="running", started_at=datetime.utcnow(), total=len(document_ids), current=0)
        
        policy_decisions = []
        
        # Policy rules
        policy_rules = [
            "block_sanctioned_parties",
            "flag_high_risk_jurisdiction",
            "check_collateral_requirements",
            "verify_esg_compliance",
            "check_credit_limits",
            "verify_kyc_compliance"
        ]
        
        for document_id in document_ids:
            try:
                document = self.db.query(Document).filter(Document.id == document_id).first()
                if not document:
                    continue
                
                # Only create for documents in workflow (submitted+)
                workflow = self.db.query(Workflow).filter(Workflow.document_id == document_id).first()
                if not workflow or workflow.state == WorkflowState.DRAFT.value:
                    continue
                
                # Select decision based on distribution
                decision = random.choices(
                    ["ALLOW", "FLAG", "BLOCK"],
                    weights=[0.70, 0.20, 0.10]
                )[0]
                
                # Select rule
                rule_applied = random.choice(policy_rules)
                
                # Generate CDM event
                import uuid
                from app.models.cdm_events import generate_cdm_policy_evaluation
                
                transaction_id = document.deal.deal_id if document.deal else f"DOC-{document_id}"
                
                policy_event = generate_cdm_policy_evaluation(
                    transaction_id=transaction_id,
                    transaction_type="facility_creation",
                    decision=decision,
                    rule_applied=rule_applied,
                    related_event_identifiers=[],
                    evaluation_trace=[{
                        "rule": rule_applied,
                        "decision": decision,
                        "document_id": document_id
                    }],
                    matched_rules=[rule_applied]
                )
                
                # Create policy decision
                policy_decision = PolicyDecision(
                    transaction_id=transaction_id,
                    transaction_type="facility_creation",
                    decision=decision,
                    rule_applied=rule_applied,
                    trace_id=str(uuid.uuid4()),
                    document_id=document_id,
                    trace=[{
                        "rule": rule_applied,
                        "decision": decision,
                        "document_id": document_id
                    }],
                    matched_rules=[rule_applied],
                    cdm_events=policy_event,
                    created_at=workflow.submitted_at or document.created_at
                )
                
                self.db.add(policy_decision)
                policy_decisions.append(policy_decision)
                self._update_status("policy_decisions", current=len(policy_decisions), progress=len(policy_decisions) / len(document_ids))
                
            except Exception as e:
                error_msg = f"Failed to create policy decision for document {document_id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("policy_decisions", errors=[error_msg])
                continue
        
        self.db.commit()
        
        self._update_status("policy_decisions", status="completed", completed_at=datetime.utcnow())
        
        return policy_decisions
    
    def _create_green_finance_policy_decisions(
        self,
        loan_assets: List[LoanAsset],
        deals: List[Deal]
    ) -> List[PolicyDecision]:
        """
        Create green finance policy decisions for demo loan assets.
        
        Args:
            loan_assets: List of LoanAsset objects
            deals: List of Deal objects (for deal_id mapping)
            
        Returns:
            List of PolicyDecision objects
        """
        if not loan_assets:
            return []
        
        policy_decisions = []
        deal_id_map = {deal.deal_id: deal.id for deal in deals}
        
        try:
            from app.services.policy_service import PolicyService
            from app.services.policy_engine_factory import get_policy_engine
            from app.models.cdm import CreditAgreement
            
            policy_service = PolicyService(get_policy_engine())
            
            for loan_asset in loan_assets:
                try:
                    if not loan_asset.geo_lat or not loan_asset.geo_lon:
                        continue
                    
                    deal_id = deal_id_map.get(loan_asset.loan_id)
                    
                    # Get deal for credit agreement context
                    deal = next((d for d in deals if d.deal_id == loan_asset.loan_id), None)
                    
                    # Create basic CreditAgreement for evaluation
                    credit_agreement = None
                    if deal:
                        credit_agreement = CreditAgreement(
                            deal_id=deal.deal_id,
                            loan_identification_number=deal.deal_id,
                            sustainability_linked=deal.deal_data.get("sustainability_linked", False) if deal.deal_data else False
                        )
                    
                    # Evaluate green finance compliance
                    green_finance_result = policy_service.evaluate_green_finance_compliance(
                        credit_agreement=credit_agreement,
                        loan_asset=loan_asset,
                        document_id=None
                    )
                    
                    # Create PolicyDecision
                    policy_decision = PolicyDecision(
                        transaction_id=loan_asset.loan_id,
                        deal_id=deal_id,
                        loan_asset_id=loan_asset.id,
                        decision=green_finance_result.decision,
                        rule_applied=green_finance_result.rule_applied,
                        trace_id=green_finance_result.trace_id,
                        policy_evaluation_trace=green_finance_result.trace,
                        matched_rules=green_finance_result.matched_rules,
                        created_at=datetime.utcnow()
                    )
                    
                    self.db.add(policy_decision)
                    policy_decisions.append(policy_decision)
                    
                except Exception as e:
                    logger.warning(f"Failed to create green finance policy decision for loan asset {loan_asset.id}: {e}")
                    continue
            
            self.db.commit()
            
        except Exception as e:
            logger.error(f"Failed to create green finance policy decisions: {e}", exc_info=True)
            self.db.rollback()
        
        return policy_decisions
    
    def _create_green_finance_assessments(
        self,
        loan_assets: List[LoanAsset],
        deals: List[Deal]
    ) -> List[GreenFinanceAssessment]:
        """
        Create green finance assessments for demo loan assets.
        
        Args:
            loan_assets: List of LoanAsset objects
            deals: List of Deal objects (for deal_id mapping)
            
        Returns:
            List of GreenFinanceAssessment objects
        """
        if not loan_assets:
            return []
        
        self._update_status("green_finance_assessments", status="running", 
                          started_at=datetime.utcnow(), total=len(loan_assets), current=0)
        
        assessments = []
        deal_id_map = {deal.deal_id: deal.id for deal in deals}
        
        for loan_asset in loan_assets:
            try:
                if not loan_asset.geo_lat or not loan_asset.geo_lon:
                    continue
                
                # Get deal_id from loan_id
                deal_id = deal_id_map.get(loan_asset.loan_id)
                
                # Extract green finance metrics
                green_metrics = loan_asset.green_finance_metrics or {}
                osm_metrics = green_metrics.get("osm_metrics", {})
                air_quality = green_metrics.get("air_quality", {})
                sustainability_components = green_metrics.get("sustainability_components", {})
                
                # Build environmental metrics
                environmental_metrics = {
                    "air_quality_index": loan_asset.air_quality_index,
                    "pm25": air_quality.get("pm25"),
                    "pm10": air_quality.get("pm10"),
                    "no2": air_quality.get("no2")
                }
                
                # Build urban activity metrics
                urban_activity_metrics = {
                    "building_count": osm_metrics.get("building_count"),
                    "building_density": osm_metrics.get("building_density"),
                    "road_density": osm_metrics.get("road_density"),
                    "green_infrastructure_coverage": osm_metrics.get("green_infrastructure_coverage")
                }
                
                # Calculate SDG alignment (simplified for demo)
                sdg_alignment = {
                    "sdg_11": min(1.0, (loan_asset.composite_sustainability_score or 0.5) * 1.1),  # Sustainable Cities
                    "sdg_13": min(1.0, (1.0 - (loan_asset.air_quality_index or 100) / 300.0)),  # Climate Action
                    "sdg_15": loan_asset.last_verified_score or 0.5,  # Life on Land (NDVI)
                    "overall_alignment": loan_asset.composite_sustainability_score or 0.5,
                    "aligned_goals": [],
                    "needs_improvement": []
                }
                
                # Determine aligned goals (>=70%) and needs improvement (<50%)
                for goal, score in [("SDG 11", sdg_alignment["sdg_11"]), 
                                   ("SDG 13", sdg_alignment["sdg_13"]), 
                                   ("SDG 15", sdg_alignment["sdg_15"])]:
                    if score >= 0.7:
                        sdg_alignment["aligned_goals"].append(goal)
                    elif score < 0.5:
                        sdg_alignment["needs_improvement"].append(goal)
                
                # Create assessment
                assessment = GreenFinanceAssessment(
                    transaction_id=loan_asset.loan_id,
                    deal_id=deal_id,
                    loan_asset_id=loan_asset.id,
                    location_lat=loan_asset.geo_lat,
                    location_lon=loan_asset.geo_lon,
                    location_type=loan_asset.location_type,
                    location_confidence=green_metrics.get("location_confidence"),
                    environmental_metrics=environmental_metrics,
                    urban_activity_metrics=urban_activity_metrics,
                    sustainability_score=loan_asset.composite_sustainability_score,
                    sustainability_components=sustainability_components,
                    sdg_alignment=sdg_alignment,
                    assessed_at=loan_asset.last_verified_at or datetime.utcnow(),
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow()
                )
                
                self.db.add(assessment)
                assessments.append(assessment)
                self._update_status("green_finance_assessments", 
                                  current=len(assessments), 
                                  progress=len(assessments) / len(loan_assets))
                
            except Exception as e:
                error_msg = f"Failed to create green finance assessment for loan asset {loan_asset.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("green_finance_assessments", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("green_finance_assessments", status="completed", 
                          completed_at=datetime.utcnow())
        
        return assessments
    
    def _create_deal_storage(self, deal_id: str) -> Path:
        """
        Create storage structure for a deal.
        
        Args:
            deal_id: Deal ID
            
        Returns:
            Path to deal storage directory
        """
        from app.services.file_storage_service import FileStorageService
        
        file_storage = FileStorageService()
        
        # Create deal folder structure
        # Note: FileStorageService.create_deal_folder expects user_id, but for demo we'll use a default
        # We'll create the demo-specific structure manually
        base_path = Path("storage/deals/demo") / deal_id
        base_path.mkdir(parents=True, exist_ok=True)
        
        # Create subdirectories
        (base_path / "documents").mkdir(exist_ok=True)
        (base_path / "extractions").mkdir(exist_ok=True)
        (base_path / "generated").mkdir(exist_ok=True)
        (base_path / "notes").mkdir(exist_ok=True)
        
        logger.debug(f"Created storage structure for deal {deal_id} at {base_path}")
        return base_path
    
    def _store_generated_documents(self, deals: List[Deal], documents: List[Document], document_versions: List[DocumentVersion], generated_docs: List[GeneratedDocument]) -> None:
        """
        Store generated documents as files (TXT, DOCX, JSON) using FileStorageService.
        
        Args:
            deals: List of Deal objects
            documents: List of Document objects
            document_versions: List of DocumentVersion objects
            generated_docs: List of GeneratedDocument objects
        """
        self._update_status("file_storage", status="running", started_at=datetime.utcnow(), total=len(document_versions) + len(generated_docs), current=0)
        
        from app.services.file_storage_service import FileStorageService
        import json
        
        file_storage = FileStorageService()
        
        # Store document versions as TXT and JSON
        for doc_version in document_versions:
            try:
                document = next((d for d in documents if d.id == doc_version.document_id), None)
                if not document or not document.deal:
                    continue
                
                deal_id = document.deal.deal_id
                storage_path = self._create_deal_storage(deal_id)
                
                # Store extracted text as TXT
                txt_file = storage_path / "documents" / f"document_{document.id}_v{doc_version.version_number}.txt"
                with open(txt_file, "w", encoding="utf-8") as f:
                    f.write(doc_version.original_text or "")  # Use original_text, not extracted_text
                
                # Store CDM JSON in extractions folder
                if doc_version.extracted_data:
                    json_file = storage_path / "extractions" / f"document_{document.id}_v{doc_version.version_number}_cdm.json"
                    with open(json_file, "w", encoding="utf-8") as f:
                        json.dump(doc_version.extracted_data, f, indent=2, default=str)
                
                self._update_status("file_storage", current=self._status.get("file_storage", SeedingStatus("file_storage")).current + 1)
                
            except Exception as e:
                error_msg = f"Failed to store document version {doc_version.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("file_storage", errors=[error_msg])
                continue
        
        # Store generated documents (already have file_path set by DocumentGenerationService)
        # Just ensure they're in the right location
        for gen_doc in generated_docs:
            try:
                if gen_doc.source_document:
                    deal = gen_doc.source_document.deal
                    if deal:
                        deal_id = deal.deal_id
                        storage_path = self._create_deal_storage(deal_id)
                        
                        # Update file path if needed
                        if gen_doc.file_path and "storage/deals/demo" not in gen_doc.file_path:
                            # Move to generated folder
                            filename = Path(gen_doc.file_path).name
                            new_path = storage_path / "generated" / filename
                            gen_doc.file_path = str(new_path)
                            
                            # If file exists at old path, move it
                            old_path = Path(gen_doc.file_path)
                            if old_path.exists():
                                new_path.parent.mkdir(parents=True, exist_ok=True)
                                import shutil
                                shutil.move(str(old_path), str(new_path))
                
                self._update_status("file_storage", current=self._status.get("file_storage", SeedingStatus("file_storage")).current + 1)
                
            except Exception as e:
                error_msg = f"Failed to store generated document {gen_doc.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                self._update_status("file_storage", errors=[error_msg])
                continue
        
        self.db.commit()
        self._update_status("file_storage", status="completed", completed_at=datetime.utcnow())
    
    def index_demo_documents(self, document_ids: List[int]) -> int:
        """
        Load documents into ChromaDB using DocumentRetrievalService with embeddings and metadata.
        
        Args:
            document_ids: List of document IDs to index
            
        Returns:
            Number of documents indexed
        """
        self._update_status("chromadb_indexing", status="running", started_at=datetime.utcnow(), total=len(document_ids), current=0)
        
        indexed_count = 0
        
        try:
            from app.chains.document_retrieval_chain import DocumentRetrievalService
            from app.core.llm_client import get_embeddings_model
            
            doc_retrieval = DocumentRetrievalService(collection_name="creditnexus_documents")
            embeddings_model = get_embeddings_model()
            
            for document_id in document_ids:
                try:
                    document = self.db.query(Document).filter(Document.id == document_id).first()
                    if not document:
                        continue
                    
                    # Get document version with extracted text
                    doc_version = None
                    if document.current_version_id:
                        doc_version = self.db.query(DocumentVersion).filter(
                            DocumentVersion.id == document.current_version_id
                        ).first()
                    
                    if not doc_version or not doc_version.original_text:
                        logger.warning(f"Document {document_id} has no original text, skipping ChromaDB indexing")
                        continue
                    
                    # Generate embedding
                    text = doc_version.original_text[:10000]  # Limit text length - use original_text, not extracted_text
                    embedding = embeddings_model.embed_query(text)
                    
                    # Prepare metadata
                    metadata = {
                        "document_id": str(document.id),
                        "deal_id": str(document.deal_id) if document.deal_id else None,
                        "borrower_name": document.borrower_name or "",
                        "is_demo": "true",
                        "created_at": document.created_at.isoformat() if document.created_at else None,
                        "title": document.title or ""
                    }
                    
                    # Add to ChromaDB
                    doc_retrieval.collection.add(
                        ids=[f"demo_doc_{document_id}"],
                        embeddings=[embedding],
                        documents=[text],
                        metadatas=[{k: str(v) if v is not None else "" for k, v in metadata.items()}]
                    )
                    
                    indexed_count += 1
                    self._update_status("chromadb_indexing", current=indexed_count, progress=indexed_count / len(document_ids))
                    
                except Exception as e:
                    error_msg = f"Failed to index document {document_id} in ChromaDB: {str(e)}"
                    logger.error(error_msg, exc_info=True)
                    self._update_status("chromadb_indexing", errors=[error_msg])
                    continue
            
            self._update_status("chromadb_indexing", status="completed", completed_at=datetime.utcnow())
            
        except ImportError:
            logger.warning("ChromaDB not available, skipping document indexing")
            self._update_status("chromadb_indexing", status="skipped", completed_at=datetime.utcnow())
        except Exception as e:
            error_msg = f"Failed to initialize ChromaDB indexing: {str(e)}"
            logger.error(error_msg, exc_info=True)
            self._update_status("chromadb_indexing", status="failed", completed_at=datetime.utcnow(), errors=[error_msg])
        
        # Create index.json in ChromaDB seed folder
        try:
            chroma_seed_path = Path("chroma_db/seeds/demo")
            chroma_seed_path.mkdir(parents=True, exist_ok=True)
            
            index_data = {
                "collection": "creditnexus_documents",
                "demo_documents": indexed_count,
                "indexed_at": datetime.utcnow().isoformat(),
                "metadata_tags": {
                    "is_demo": "true",
                    "source": "demo_data_service",
                    "document_count": indexed_count
                },
                "document_ids": [str(doc_id) for doc_id in document_ids[:indexed_count]]
            }
            
            index_file = chroma_seed_path / "index.json"
            with open(index_file, "w", encoding="utf-8") as f:
                json.dump(index_data, f, indent=2, default=str)
            
            logger.info(f"Created ChromaDB seed index at {index_file}")
            
        except Exception as e:
            logger.warning(f"Failed to create ChromaDB seed index: {e}")
        
        return indexed_count
    
    def complete_partially_filled_data(
        self,
        complete_deals: bool = True,
        complete_loan_assets: bool = True,
        complete_applications: bool = True,
        complete_documents: bool = True
    ) -> Dict[str, Any]:
        """
        Complete missing fields in partially filled synthetic data points.
        
        This method identifies demo records that have some fields filled but others are missing,
        and completes them using the same generation logic as new records.
        
        Args:
            complete_deals: Whether to complete partially filled deals
            complete_loan_assets: Whether to complete partially filled loan assets
            complete_applications: Whether to complete partially filled applications
            complete_documents: Whether to complete partially filled documents
            
        Returns:
            Dictionary with completion statistics
        """
        completion_stats = {
            "deals_completed": 0,
            "loan_assets_completed": 0,
            "applications_completed": 0,
            "documents_completed": 0,
            "errors": []
        }
        
        self._update_status("data_completion", status="running", started_at=datetime.utcnow(), total=4, current=0)
        
        try:
            # Step 1: Complete partially filled deals
            if complete_deals:
                deals_completed = self._complete_partial_deals()
                completion_stats["deals_completed"] = deals_completed
                self._update_status("data_completion", current=1, progress=0.25)
            
            # Step 2: Complete partially filled loan assets
            if complete_loan_assets:
                loan_assets_completed = self._complete_partial_loan_assets()
                completion_stats["loan_assets_completed"] = loan_assets_completed
                self._update_status("data_completion", current=2, progress=0.50)
            
            # Step 3: Complete partially filled applications
            if complete_applications:
                applications_completed = self._complete_partial_applications()
                completion_stats["applications_completed"] = applications_completed
                self._update_status("data_completion", current=3, progress=0.75)
            
            # Step 4: Complete partially filled documents
            if complete_documents:
                documents_completed = self._complete_partial_documents()
                completion_stats["documents_completed"] = documents_completed
                self._update_status("data_completion", current=4, progress=1.0)
            
            self.db.commit()
            self._update_status("data_completion", status="completed", completed_at=datetime.utcnow())
            
        except Exception as e:
            error_msg = f"Failed to complete partially filled data: {str(e)}"
            logger.error(error_msg, exc_info=True)
            completion_stats["errors"].append(error_msg)
            self._update_status("data_completion", status="failed", errors=[error_msg])
            self.db.rollback()
        
        return completion_stats
    
    def _complete_partial_deals(self) -> int:
        """
        Complete missing fields in partially filled demo deals.
        
        Returns:
            Number of deals completed
        """
        # Find demo deals with missing or incomplete deal_data
        partial_deals = self.db.query(Deal).filter(
            Deal.is_demo == True
        ).all()
        
        completed_count = 0
        
        for deal in partial_deals:
            try:
                needs_completion = False
                deal_data = deal.deal_data or {}
                
                # Check for missing critical fields in deal_data
                required_fields = [
                    "loan_amount", "term_years", "interest_rate", "industry",
                    "sustainability_linked", "esg_targets", "collateral_type"
                ]
                
                for field in required_fields:
                    if field not in deal_data or deal_data[field] is None:
                        needs_completion = True
                        break
                
                # Check if deal_type is missing
                if not deal.deal_type:
                    needs_completion = True
                
                if needs_completion:
                    # Get application data if available
                    application = deal.application
                    if application and application.application_data:
                        app_data = application.application_data
                        loan_amount = app_data.get("loan_amount", 1000000)
                        industry = app_data.get("industry", "Technology")
                    else:
                        loan_amount = deal_data.get("loan_amount", 1000000)
                        industry = deal_data.get("industry", "Technology")
                    
                    # Get industry-specific config
                    from app.prompts.demo.deal_generation import get_industry_config
                    industry_config = get_industry_config(industry)
                    
                    # Complete missing deal_data fields
                    if "loan_amount" not in deal_data or deal_data["loan_amount"] is None:
                        deal_data["loan_amount"] = loan_amount
                    
                    if "term_years" not in deal_data or deal_data["term_years"] is None:
                        term_min, term_max = industry_config["term_range"]
                        deal_data["term_years"] = random.randint(term_min, term_max)
                    
                    if "interest_rate" not in deal_data or deal_data["interest_rate"] is None:
                        rate_min, rate_max = industry_config["interest_rate_range"]
                        deal_data["interest_rate"] = round(random.uniform(rate_min, rate_max), 2)
                    
                    if "industry" not in deal_data or deal_data["industry"] is None:
                        deal_data["industry"] = industry
                    
                    if "sustainability_linked" not in deal_data or deal_data["sustainability_linked"] is None:
                        # 30% of deals are sustainability-linked
                        deal_data["sustainability_linked"] = random.random() < 0.30
                    
                    if "collateral_type" not in deal_data or deal_data["collateral_type"] is None:
                        deal_data["collateral_type"] = random.choice(industry_config["collateral_types"])
                    
                    # Complete ESG targets if sustainability-linked
                    if deal_data.get("sustainability_linked", False):
                        if "esg_targets" not in deal_data or deal_data["esg_targets"] is None:
                            deal_data["esg_targets"] = {
                                "ndvi_threshold": round(random.uniform(0.70, 0.85), 2),
                                "penalty_bps": random.randint(25, 50),
                                "measurement_frequency": "Quarterly"
                            }
                    
                    # Set deal_type if missing
                    if not deal.deal_type:
                        deal.deal_type = DealType.LOAN_APPLICATION.value
                    
                    # Update deal
                    deal.deal_data = deal_data
                    deal.updated_at = datetime.utcnow()
                    self.db.add(deal)
                    completed_count += 1
                    
            except Exception as e:
                error_msg = f"Failed to complete deal {deal.deal_id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                continue
        
        self.db.flush()
        logger.info(f"Completed {completed_count} partially filled deals")
        return completed_count
    
    def _complete_partial_loan_assets(self) -> int:
        """
        Complete missing fields in partially filled demo loan assets.
        
        Returns:
            Number of loan assets completed
        """
        # Find demo loan assets (those associated with demo deals)
        demo_deal_ids = [d.deal_id for d in self.db.query(Deal).filter(Deal.is_demo == True).all()]
        
        if not demo_deal_ids:
            return 0
        
        partial_loan_assets = self.db.query(LoanAsset).filter(
            LoanAsset.loan_id.in_(demo_deal_ids)
        ).all()
        
        completed_count = 0
        
        for loan_asset in partial_loan_assets:
            try:
                needs_completion = False
                
                # Check for missing critical fields
                if not loan_asset.original_text:
                    needs_completion = True
                if not loan_asset.legal_vector:
                    needs_completion = True
                if not loan_asset.collateral_address:
                    needs_completion = True
                if loan_asset.geo_lat is None or loan_asset.geo_lon is None:
                    needs_completion = True
                if not loan_asset.satellite_snapshot_url:
                    needs_completion = True
                if not loan_asset.geo_vector:
                    needs_completion = True
                if not loan_asset.spt_data:
                    needs_completion = True
                if not loan_asset.green_finance_metrics:
                    needs_completion = True
                if loan_asset.location_type is None:
                    needs_completion = True
                if loan_asset.air_quality_index is None:
                    needs_completion = True
                if loan_asset.composite_sustainability_score is None:
                    needs_completion = True
                
                if needs_completion:
                    # Get associated deal
                    deal = self.db.query(Deal).filter(Deal.deal_id == loan_asset.loan_id).first()
                    if not deal:
                        continue
                    
                    # Use the same generation logic as _generate_loan_assets_for_deals
                    # Get ESG targets from deal
                    esg_targets = deal.deal_data.get("esg_targets", {}) if deal.deal_data else {}
                    ndvi_threshold = esg_targets.get("ndvi_threshold", 0.75)
                    
                    # Generate or complete missing address
                    if not loan_asset.collateral_address:
                        addresses = [
                            "123 Industrial Way, Detroit, MI 48201",
                            "456 Farm Road, Napa, CA 94558",
                            "789 Agricultural Blvd, Fresno, CA 93721",
                            "321 Green Valley Lane, Austin, TX 78701"
                        ]
                        loan_asset.collateral_address = random.choice(addresses)
                    
                    # Generate or complete coordinates
                    if loan_asset.geo_lat is None or loan_asset.geo_lon is None:
                        geo_coords = {
                            "123 Industrial Way, Detroit, MI 48201": (42.3314, -83.0458),
                            "456 Farm Road, Napa, CA 94558": (38.2975, -122.2869),
                            "789 Agricultural Blvd, Fresno, CA 93721": (36.7378, -119.7871),
                            "321 Green Valley Lane, Austin, TX 78701": (30.2672, -97.7431)
                        }
                        geo_lat, geo_lon = geo_coords.get(loan_asset.collateral_address, (40.7128, -74.0060))
                        loan_asset.geo_lat = geo_lat
                        loan_asset.geo_lon = geo_lon
                    
                    # Generate or complete NDVI score
                    if loan_asset.last_verified_score is None:
                        loan_asset.last_verified_score = random.uniform(0.50, 0.95)
                    
                    # Calculate or complete risk status
                    if loan_asset.risk_status == "PENDING" or not loan_asset.risk_status:
                        if loan_asset.last_verified_score >= ndvi_threshold:
                            loan_asset.risk_status = "COMPLIANT"
                        elif loan_asset.last_verified_score >= ndvi_threshold * 0.9:
                            loan_asset.risk_status = "WARNING"
                        else:
                            loan_asset.risk_status = "BREACH"
                    
                    # Get or complete interest rates
                    base_interest_rate = deal.deal_data.get("interest_rate", 5.25) if deal.deal_data else 5.25
                    penalty_bps = esg_targets.get("penalty_bps", 25) if isinstance(esg_targets, dict) else 25
                    
                    if loan_asset.base_interest_rate is None:
                        loan_asset.base_interest_rate = float(base_interest_rate)
                    
                    if loan_asset.current_interest_rate is None:
                        current_rate = base_interest_rate + (penalty_bps / 10000 if loan_asset.risk_status == "BREACH" else 0)
                        loan_asset.current_interest_rate = float(current_rate)
                    
                    if loan_asset.penalty_bps is None:
                        loan_asset.penalty_bps = penalty_bps if loan_asset.risk_status == "BREACH" else 0
                    
                    # Generate or complete original_text
                    if not loan_asset.original_text:
                        loan_amount = deal.deal_data.get('loan_amount', 1000000) if deal.deal_data else 1000000
                        term_years = deal.deal_data.get('term_years', 5) if deal.deal_data else 5
                        loan_asset.original_text = f"""Credit Agreement for {deal.deal_id}
                        
This sustainability-linked credit agreement establishes a loan facility with the following terms:
- Loan Amount: ${loan_amount:,.0f}
- Interest Rate: {base_interest_rate}% base rate
- Term: {term_years} years
- Collateral: {loan_asset.collateral_address}
- ESG Target: NDVI threshold of {ndvi_threshold:.2f}
- Penalty: {penalty_bps} basis points for non-compliance

The borrower agrees to maintain the specified NDVI threshold and will be subject to quarterly verification."""
                    
                    # Generate or complete legal_vector
                    if not loan_asset.legal_vector:
                        loan_asset.legal_vector = [random.uniform(-1.0, 1.0) for _ in range(1536)]
                    
                    # Generate or complete satellite_snapshot_url
                    if not loan_asset.satellite_snapshot_url:
                        loan_asset.satellite_snapshot_url = f"https://josephrp.github.io/creditnexus/snapshots/{deal.deal_id}/{datetime.utcnow().strftime('%Y%m%d')}.tif"
                    
                    # Generate or complete geo_vector
                    if not loan_asset.geo_vector:
                        loan_asset.geo_vector = [random.uniform(-1.0, 1.0) for _ in range(512)]
                    
                    # Generate or complete spt_data
                    if not loan_asset.spt_data:
                        loan_asset.spt_data = {
                            "target_type": "NDVI",
                            "threshold": ndvi_threshold,
                            "measurement_frequency": "Quarterly",
                            "penalty_bps": penalty_bps,
                            "cdm_compliant": True,
                            "measurement_history": [
                                {
                                    "date": (datetime.utcnow() - timedelta(days=90*i)).isoformat(),
                                    "value": max(0.0, min(1.0, loan_asset.last_verified_score + random.uniform(-0.1, 0.1))),
                                    "verified_by": "Sentinel-2B",
                                    "confidence": random.uniform(0.85, 0.98)
                                }
                                for i in range(4)
                            ],
                            "next_verification_date": (datetime.utcnow() + timedelta(days=90)).isoformat()
                        }
                    
                    # Set spt_threshold if missing
                    if loan_asset.spt_threshold is None:
                        loan_asset.spt_threshold = ndvi_threshold
                    
                    # Generate or complete green finance metrics
                    if not loan_asset.green_finance_metrics or loan_asset.location_type is None:
                        # Location classification
                        location_type_map = {
                            "123 Industrial Way, Detroit, MI 48201": "urban",
                            "456 Farm Road, Napa, CA 94558": "rural",
                            "789 Agricultural Blvd, Fresno, CA 93721": "suburban",
                            "321 Green Valley Lane, Austin, TX 78701": "suburban"
                        }
                        location_type = location_type_map.get(loan_asset.collateral_address, random.choice(["urban", "suburban", "rural"]))
                        loan_asset.location_type = location_type
                        
                        # Generate air quality data
                        if location_type == "urban":
                            aqi = random.uniform(80, 150)
                        elif location_type == "suburban":
                            aqi = random.uniform(50, 100)
                        else:
                            aqi = random.uniform(20, 60)
                        
                        loan_asset.air_quality_index = aqi
                        
                        # Calculate composite sustainability score
                        vegetation_health = loan_asset.last_verified_score
                        if aqi <= 50:
                            air_quality_score = 1.0
                        elif aqi <= 100:
                            air_quality_score = 0.8
                        elif aqi <= 150:
                            air_quality_score = 0.6
                        elif aqi <= 200:
                            air_quality_score = 0.4
                        else:
                            air_quality_score = 0.2
                        
                        composite_score = (
                            vegetation_health * 0.30 +
                            air_quality_score * 0.25 +
                            0.45  # Simplified calculation
                        )
                        loan_asset.composite_sustainability_score = composite_score
                        
                        # Build green_finance_metrics
                        loan_asset.green_finance_metrics = {
                            "location_type": location_type,
                            "location_confidence": random.uniform(0.75, 0.95),
                            "air_quality": {
                                "pm25": random.uniform(5, 55),
                                "pm10": random.uniform(10, 100),
                                "no2": random.uniform(5, 80),
                                "data_source": "synthetic_demo"
                            },
                            "sustainability_components": {
                                "vegetation_health": vegetation_health,
                                "air_quality": air_quality_score
                            }
                        }
                    
                    # Set last_verified_at if missing
                    if not loan_asset.last_verified_at:
                        loan_asset.last_verified_at = datetime.utcnow() - timedelta(days=random.randint(1, 30))
                    
                    # Generate or complete asset_metadata
                    if not loan_asset.asset_metadata:
                        loan_asset.asset_metadata = {
                            "verification_method": "Sentinel-2B",
                            "cloud_cover": random.uniform(0.01, 0.10),
                            "classification": random.choice(["AnnualCrop", "Forest", "PermanentCrop"]),
                            "confidence": random.uniform(0.85, 0.98)
                        }
                    
                    self.db.add(loan_asset)
                    completed_count += 1
                    
            except Exception as e:
                error_msg = f"Failed to complete loan asset {loan_asset.loan_id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                continue
        
        self.db.flush()
        logger.info(f"Completed {completed_count} partially filled loan assets")
        return completed_count
    
    def _complete_partial_applications(self) -> int:
        """
        Complete missing fields in partially filled demo applications.
        
        Returns:
            Number of applications completed
        """
        # Find demo applications (those associated with demo deals)
        demo_deal_ids = [d.deal_id for d in self.db.query(Deal).filter(Deal.is_demo == True).all()]
        demo_application_ids = [
            d.application_id for d in self.db.query(Deal).filter(
                Deal.is_demo == True,
                Deal.application_id.isnot(None)
            ).all()
        ]
        
        if not demo_application_ids:
            return 0
        
        partial_applications = self.db.query(Application).filter(
            Application.id.in_(demo_application_ids)
        ).all()
        
        completed_count = 0
        
        for application in partial_applications:
            try:
                needs_completion = False
                
                # Check for missing critical fields
                if not application.application_data:
                    needs_completion = True
                elif not application.application_data.get("loan_amount"):
                    needs_completion = True
                elif not application.application_data.get("industry"):
                    needs_completion = True
                
                if application.application_type == ApplicationType.BUSINESS.value:
                    if not application.business_data:
                        needs_completion = True
                elif application.application_type == ApplicationType.INDIVIDUAL.value:
                    if not application.individual_data:
                        needs_completion = True
                
                if needs_completion:
                    # Complete application_data
                    if not application.application_data:
                        application.application_data = {}
                    
                    app_data = application.application_data
                    
                    if not app_data.get("loan_amount"):
                        app_data["loan_amount"] = random.randint(100000, 10000000)
                    
                    if not app_data.get("industry"):
                        from app.prompts.demo.deal_generation import get_industry_weights
                        industry_weights = get_industry_weights()
                        industries = list(industry_weights.keys())
                        weights = list(industry_weights.values())
                        app_data["industry"] = random.choices(industries, weights=weights)[0]
                    
                    if not app_data.get("purpose"):
                        app_data["purpose"] = random.choice(["Working capital", "Expansion", "Refinancing", "Equipment purchase", "Inventory"])
                    
                    if not app_data.get("years_in_business"):
                        app_data["years_in_business"] = random.randint(2, 50)
                    
                    if not app_data.get("annual_revenue"):
                        app_data["annual_revenue"] = app_data["loan_amount"] * random.uniform(2, 10)
                    
                    if not app_data.get("credit_score"):
                        app_data["credit_score"] = random.randint(650, 850)
                    
                    application.application_data = app_data
                    
                    # Complete business_data for business applications
                    if application.application_type == ApplicationType.BUSINESS.value:
                        if not application.business_data:
                            application.business_data = {}
                        
                        business_data = application.business_data
                        
                        if not business_data.get("company_name"):
                            business_data["company_name"] = f"Demo Company {application.id}"
                        
                        if not business_data.get("tax_id"):
                            business_data["tax_id"] = f"{random.randint(10, 99)}-{random.randint(1000000, 9999999)}"
                        
                        if not business_data.get("legal_structure"):
                            business_data["legal_structure"] = random.choice(["Corporation", "LLC", "Partnership"])
                        
                        if not business_data.get("number_of_employees"):
                            business_data["number_of_employees"] = random.randint(10, 5000)
                        
                        if not business_data.get("business_address"):
                            business_data["business_address"] = f"{random.randint(1, 9999)} Business St, City, ST {random.randint(10000, 99999)}"
                        
                        application.business_data = business_data
                    
                    # Complete individual_data for individual applications
                    elif application.application_type == ApplicationType.INDIVIDUAL.value:
                        if not application.individual_data:
                            application.individual_data = {}
                        
                        individual_data = application.individual_data
                        
                        if not individual_data.get("first_name"):
                            individual_data["first_name"] = "John"
                        
                        if not individual_data.get("last_name"):
                            individual_data["last_name"] = f"Doe{application.id}"
                        
                        if not individual_data.get("date_of_birth"):
                            individual_data["date_of_birth"] = (datetime.utcnow() - timedelta(days=random.randint(25*365, 65*365))).date().isoformat()
                        
                        if not individual_data.get("ssn_last_4"):
                            individual_data["ssn_last_4"] = f"{random.randint(1000, 9999)}"
                        
                        if not individual_data.get("employment_status"):
                            individual_data["employment_status"] = random.choice(["Employed", "Self-employed", "Retired"])
                        
                        if not individual_data.get("annual_income"):
                            individual_data["annual_income"] = app_data.get("annual_revenue", 100000)
                        
                        if not individual_data.get("residential_address"):
                            individual_data["residential_address"] = f"{random.randint(1, 9999)} Residential Ave, City, ST {random.randint(10000, 99999)}"
                        
                        if not individual_data.get("years_at_address"):
                            individual_data["years_at_address"] = random.randint(1, 20)
                        
                        application.individual_data = individual_data
                    
                    self.db.add(application)
                    completed_count += 1
                    
            except Exception as e:
                error_msg = f"Failed to complete application {application.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                continue
        
        self.db.flush()
        logger.info(f"Completed {completed_count} partially filled applications")
        return completed_count
    
    def _complete_partial_documents(self) -> int:
        """
        Complete missing fields in partially filled demo documents.
        
        Returns:
            Number of documents completed
        """
        # Find demo documents (those associated with demo deals)
        demo_deal_ids = [d.id for d in self.db.query(Deal).filter(Deal.is_demo == True).all()]
        
        if not demo_deal_ids:
            return 0
        
        partial_documents = self.db.query(Document).filter(
            Document.deal_id.in_(demo_deal_ids)
        ).all()
        
        completed_count = 0
        
        for document in partial_documents:
            try:
                needs_completion = False
                
                # Check for missing critical fields
                if not document.title:
                    needs_completion = True
                if not document.document_type:
                    needs_completion = True
                if not document.status:
                    needs_completion = True
                
                if needs_completion:
                    # Complete title if missing
                    if not document.title:
                        document.title = f"Document for Deal {document.deal_id}"
                    
                    # Complete document_type if missing
                    if not document.document_type:
                        document.document_type = random.choice([
                            "credit_agreement", "term_sheet", "amendment", "notice", "certificate"
                        ])
                    
                    # Complete status if missing
                    if not document.status:
                        document.status = "draft"
                    
                    self.db.add(document)
                    completed_count += 1
                    
            except Exception as e:
                error_msg = f"Failed to complete document {document.id}: {str(e)}"
                logger.error(error_msg, exc_info=True)
                continue
        
        self.db.flush()
        logger.info(f"Completed {completed_count} partially filled documents")
        return completed_count