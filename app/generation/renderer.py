"""
Word document renderer for LMA template generation.

Handles placeholder replacement in Word documents including paragraphs,
tables, headers, and footers.
"""

import logging
import re
import json
from pathlib import Path
from typing import Dict, Optional, Any, List
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.generation.field_parser import FieldPathParser
from app.models.cdm import CreditAgreement

logger = logging.getLogger(__name__)


class DocumentRenderer:
    """
    Renders Word documents by replacing placeholders with field values.
    
    Supports placeholder replacement in:
    - Paragraphs
    - Tables
    - Headers and footers
    - Content controls (structured document tags)
    - Text boxes and shapes
    """
    
    # Pattern to match placeholders like [BORROWER_NAME] or [COMMITMENT_AMOUNT]
    PLACEHOLDER_PATTERN = re.compile(r'\[([A-Z_][A-Z0-9_]*)\]')
    # Pattern to match curly brace placeholders like {{representations_and_warranties}}
    CURLY_PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
    # Pattern to match single curly brace placeholders like {facilities[0].facility_name}
    SINGLE_CURLY_PLACEHOLDER_PATTERN = re.compile(r'\{([^}]+)\}')
    
    def __init__(self):
        """Initialize document renderer."""
        self.parser = FieldPathParser()
        logger.debug("DocumentRenderer initialized")
    
    def render_template(
        self,
        template_path: str,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> DocumentType:
        """
        Render template by replacing placeholders with field values.
        
        Args:
            template_path: Path to template Word document
            field_values: Dictionary mapping placeholder names to values
                Example: {"[BORROWER_NAME]": "ACME Corp", "[COMMITMENT_AMOUNT]": "$10,000,000.00 USD"}
            
        Returns:
            Rendered Document instance
            
        Raises:
            FileNotFoundError: If template file doesn't exist
            IOError: If template cannot be read
        """
        template_file = Path(template_path)
        if not template_file.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            # Load template document
            doc = Document(str(template_file))
            
            # Replace placeholders in all document elements
            self._replace_placeholders(doc, field_values, cdm_data)
            self._replace_in_tables(doc, field_values, cdm_data)
            self._replace_in_headers_footers(doc, field_values, cdm_data)
            self._replace_in_content_controls(doc, field_values, cdm_data)
            self._replace_in_textboxes_shapes(doc, field_values, cdm_data)
            
            logger.info(f"Rendered template {template_path} with {len(field_values)} field(s)")
            return doc
            
        except Exception as e:
            raise IOError(f"Failed to render template {template_path}: {e}") from e
    
    def _replace_placeholders(
        self,
        doc: DocumentType,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> None:
        """
        Replace placeholders in document paragraphs.
        
        Args:
            doc: Document instance
            field_values: Dictionary of placeholder to value mappings
        """
        for idx, paragraph in enumerate(doc.paragraphs):
            self._replace_in_paragraph(paragraph, field_values, cdm_data)
    
    def _replace_in_paragraph(
        self,
        paragraph: Paragraph,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> int:
        """
        Replace placeholders in a single paragraph.
        
        Args:
            paragraph: Paragraph instance
            field_values: Dictionary of placeholder to value mappings
            
        Returns:
            Number of replacements made
        """
        if not paragraph.text:
            return 0
        
        # Check for alternative placeholder formats
        curly_placeholders = self.CURLY_PLACEHOLDER_PATTERN.findall(paragraph.text)
        # Find all double curly matches with their positions to exclude from single curly matching
        double_curly_ranges = []
        for match in self.CURLY_PLACEHOLDER_PATTERN.finditer(paragraph.text):
            double_curly_ranges.append((match.start(), match.end()))
        
        # Find single curly placeholders, but exclude those inside double curly braces
        single_curly_placeholders = []
        for match in self.SINGLE_CURLY_PLACEHOLDER_PATTERN.finditer(paragraph.text):
            # Check if this single curly match is inside any double curly match
            is_inside_double = any(
                double_start <= match.start() < double_end or double_start < match.end() <= double_end
                for double_start, double_end in double_curly_ranges
            )
            if not is_inside_double:
                single_curly_placeholders.append(match.group(1))
        
        bracket_placeholders = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
        
        # Build replacement text
        text = paragraph.text
        replacements_made = 0
        
        # First, handle bracket format placeholders [FIELD_NAME]
        for placeholder_name in bracket_placeholders:
            placeholder = f"[{placeholder_name}]"
            
            if placeholder in field_values:
                value = str(field_values[placeholder])
                text = text.replace(placeholder, value)
                replacements_made += 1
            else:
                logger.debug(f"Placeholder {placeholder} not found in field_values, leaving as-is")
        
        # Second, handle double curly brace format placeholders {{field_name}}
        # Map curly format to bracket format for field lookup
        for curly_placeholder in curly_placeholders:
            # Normalize: convert to uppercase with underscores, remove brackets/quotes
            # e.g., "parties[role=\"Borrower\"].name" -> "BORROWER_NAME"
            # e.g., "representations_and_warranties" -> "REPRESENTATIONS_AND_WARRANTIES"
            normalized = self._normalize_curly_placeholder(curly_placeholder)
            bracket_key = f"[{normalized}]"
            curly_full = f"{{{{{curly_placeholder}}}}}"
            
            if bracket_key in field_values:
                value = str(field_values[bracket_key])
                text = text.replace(curly_full, value)
                replacements_made += 1
            else:
                # Try to evaluate as CDM path if cdm_data is available
                value = None
                if cdm_data:
                    try:
                        cdm_value = self.parser.get_nested_value(cdm_data, curly_placeholder)
                        if cdm_value is not None:
                            # Format the value
                            from decimal import Decimal
                            from datetime import date
                            if hasattr(cdm_value, 'value'):  # Enum
                                value = str(cdm_value.value)
                            elif isinstance(cdm_value, (int, float, Decimal)):
                                value = str(cdm_value)
                            elif isinstance(cdm_value, date):
                                value = cdm_value.isoformat()
                            else:
                                value = str(cdm_value)
                    except Exception as e:
                        logger.debug(f"Error evaluating CDM path {curly_placeholder}: {e}")
                
                if value:
                    text = text.replace(curly_full, value)
                    replacements_made += 1
                else:
                    logger.debug(f"Curly placeholder {curly_full} (normalized to {bracket_key}) not found in field_values or CDM data")
        
        # Third, handle single curly brace format placeholders {field_name}
        # These are typically CDM field paths that need to be mapped
        for single_curly in single_curly_placeholders:
            # Skip if already handled as double curly
            if single_curly in curly_placeholders:
                continue
            
            # Normalize single curly placeholder
            normalized = self._normalize_curly_placeholder(single_curly)
            bracket_key = f"[{normalized}]"
            single_curly_full = f"{{{single_curly}}}"
            
            value = None
            
            if bracket_key in field_values:
                value = str(field_values[bracket_key])
            elif cdm_data:
                # Try to evaluate as CDM path
                try:
                    cdm_value = self.parser.get_nested_value(cdm_data, single_curly)
                    if cdm_value is not None:
                        # Format the value
                        from decimal import Decimal
                        from datetime import date
                        if hasattr(cdm_value, 'value'):  # Enum
                            value = str(cdm_value.value)
                        elif isinstance(cdm_value, (int, float, Decimal)):
                            value = str(cdm_value)
                        elif isinstance(cdm_value, date):
                            value = cdm_value.isoformat()
                        else:
                            value = str(cdm_value)
                except Exception as e:
                    logger.debug(f"Error evaluating CDM path {single_curly}: {e}")
            
            if value:
                text = text.replace(single_curly_full, value)
                replacements_made += 1
            else:
                logger.debug(f"Single curly placeholder {single_curly_full} (normalized to {bracket_key}) not found in field_values or CDM data")
        
        # Replace paragraph text
        # Clear existing runs and add new text
        if replacements_made > 0:
            paragraph.clear()
            paragraph.add_run(text)
        
        return replacements_made
    
    def _normalize_curly_placeholder(self, curly_placeholder: str) -> str:
        """
        Normalize curly brace placeholder to bracket format key.
        
        Examples:
        - "parties[role=\"Borrower\"].name" -> "BORROWER_NAME"
        - "representations_and_warranties" -> "REPRESENTATIONS_AND_WARRANTIES"
        - "facilities[0].facility_name" -> "FACILITY_NAME"
        - "conditions_precedent" -> "CONDITIONS_PRECEDENT"
        - "covenants" -> "COVENANTS"
        - "events_of_default" -> "EVENTS_OF_DEFAULT"
        
        Args:
            curly_placeholder: Placeholder text from {{...}}
            
        Returns:
            Normalized uppercase key for field_values lookup
        """
        # Handle simple field names (already normalized, just uppercase)
        if not '[' in curly_placeholder and not '.' in curly_placeholder:
            return curly_placeholder.upper().replace('-', '_')
        
        # Extract field name from path expressions
        # e.g., "parties[role=\"Borrower\"].name" -> extract "Borrower" and "name"
        role_match = re.search(r'role\s*=\s*["\']?(\w+)["\']?', curly_placeholder)
        if role_match:
            role = role_match.group(1)
            # Get the field after the dot
            field_match = re.search(r'\.(\w+)(?:\.|$|})', curly_placeholder)
            if field_match:
                field = field_match.group(1)
                # Combine: "Borrower" + "name" -> "BORROWER_NAME"
                return f"{role.upper()}_{field.upper()}"
        
        # Handle array access: facilities[0].facility_name -> FACILITY_NAME
        # Remove array indices: facilities[0] -> facilities
        text = re.sub(r'\[\d+\]', '', curly_placeholder)
        
        # Extract last part of path
        last_part = text.split('.')[-1] if '.' in text else text
        
        # Convert snake_case or camelCase to UPPER_SNAKE_CASE
        # Split on dots, underscores, or camelCase boundaries
        parts = re.split(r'[._]|(?<=[a-z])(?=[A-Z])', last_part)
        normalized = '_'.join(p.upper() for p in parts if p)
        
        # Clean up: remove any remaining brackets, quotes, spaces
        normalized = re.sub(r'[^A-Z0-9_]', '', normalized)
        
        return normalized
    
    def _replace_in_tables(
        self,
        doc: DocumentType,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> None:
        """
        Replace placeholders in document tables.
        
        Args:
            doc: Document instance
            field_values: Dictionary of placeholder to value mappings
        """
        # Iterate through all tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Replace in each cell's paragraphs
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, field_values)
    
    def _replace_in_headers_footers(
        self,
        doc: DocumentType,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> None:
        """
        Replace placeholders in document headers and footers.
        
        Args:
            doc: Document instance
            field_values: Dictionary of placeholder to value mappings
        """
        # Process each section's headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_in_paragraph(paragraph, field_values)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_in_paragraph(paragraph, field_values)
            
            # First page header (if different)
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    self._replace_in_paragraph(paragraph, field_values)
            
            # First page footer (if different)
            if section.first_page_footer:
                for paragraph in section.first_page_footer.paragraphs:
                    self._replace_in_paragraph(paragraph, field_values)
    
    def save_document(
        self,
        doc: DocumentType,
        output_path: str
    ) -> str:
        """
        Save rendered document to file.
        
        Args:
            doc: Document instance
            output_path: Path to save the document
            
        Returns:
            Absolute path to saved file
        """
        output_file = Path(output_path)
        
        # Create parent directory if needed
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_file))
        
        logger.info(f"Saved rendered document: {output_file} ({output_file.stat().st_size} bytes)")
        return str(output_file.absolute())
    
    def export_to_pdf(
        self,
        doc: DocumentType,
        output_path: str
    ) -> str:
        """
        Export document to PDF format.
        
        Note: This requires docx2pdf or similar library. For now, this is a placeholder.
        In production, you would use:
        - docx2pdf (Windows only, requires Microsoft Word)
        - LibreOffice headless (cross-platform)
        - Cloud conversion service
        
        Args:
            doc: Document instance
            output_path: Path to save the PDF
            
        Returns:
            Absolute path to saved PDF file
            
        Raises:
            NotImplementedError: PDF export not yet implemented
        """
        # For now, raise NotImplementedError
        # In production, implement using one of:
        # 1. docx2pdf (Windows only)
        # 2. LibreOffice headless conversion
        # 3. Cloud service (e.g., CloudConvert API)
        
        raise NotImplementedError(
            "PDF export not yet implemented. "
            "Options: 1) Install docx2pdf (Windows only), "
            "2) Use LibreOffice headless, "
            "3) Use cloud conversion service"
        )
        
        # Example implementation with docx2pdf (if available):
        # try:
        #     from docx2pdf import convert
        #     pdf_path = output_path.replace('.docx', '.pdf')
        #     convert(str(doc_path), pdf_path)
        #     return pdf_path
        # except ImportError:
        #     raise NotImplementedError("docx2pdf not installed")
    
    def validate_rendered_document(
        self,
        doc: DocumentType,
        expected_placeholders: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Validate that all placeholders were replaced in the rendered document.
        
        Args:
            doc: Rendered Document instance
            expected_placeholders: Optional list of placeholder names that should have been replaced
            
        Returns:
            Dictionary with validation results:
            - valid: bool - True if no placeholders remain
            - remaining_placeholders: List[str] - List of placeholder names still present
            - placeholder_locations: Dict[str, List[str]] - Mapping of placeholder to locations
            - total_placeholders_found: int - Total count of remaining placeholders
        """
        remaining_placeholders = []
        placeholder_locations = {}
        
        # Check paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            # Check for bracket placeholders [FIELD_NAME]
            for match in self.PLACEHOLDER_PATTERN.finditer(text):
                placeholder = match.group(0)
                if placeholder not in placeholder_locations:
                    placeholder_locations[placeholder] = []
                placeholder_locations[placeholder].append(f"paragraph_{para_idx}")
                if placeholder not in remaining_placeholders:
                    remaining_placeholders.append(placeholder)
            
            # Check for double curly placeholders {{field}}
            for match in self.CURLY_PLACEHOLDER_PATTERN.finditer(text):
                placeholder = match.group(0)
                if placeholder not in placeholder_locations:
                    placeholder_locations[placeholder] = []
                placeholder_locations[placeholder].append(f"paragraph_{para_idx}")
                if placeholder not in remaining_placeholders:
                    remaining_placeholders.append(placeholder)
            
            # Check for single curly placeholders {field}
            for match in self.SINGLE_CURLY_PLACEHOLDER_PATTERN.finditer(text):
                # Filter out matches that are inside double curly braces
                is_inside_double = False
                for dc_match in self.CURLY_PLACEHOLDER_PATTERN.finditer(text):
                    if dc_match.start() <= match.start() and match.end() <= dc_match.end():
                        is_inside_double = True
                        break
                if not is_inside_double:
                    placeholder = match.group(0)
                    if placeholder not in placeholder_locations:
                        placeholder_locations[placeholder] = []
                    placeholder_locations[placeholder].append(f"paragraph_{para_idx}")
                    if placeholder not in remaining_placeholders:
                        remaining_placeholders.append(placeholder)
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    # Check for bracket placeholders
                    for match in self.PLACEHOLDER_PATTERN.finditer(text):
                        placeholder = match.group(0)
                        if placeholder not in placeholder_locations:
                            placeholder_locations[placeholder] = []
                        placeholder_locations[placeholder].append(f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
                        if placeholder not in remaining_placeholders:
                            remaining_placeholders.append(placeholder)
                    
                    # Check for curly placeholders
                    for match in self.CURLY_PLACEHOLDER_PATTERN.finditer(text):
                        placeholder = match.group(0)
                        if placeholder not in placeholder_locations:
                            placeholder_locations[placeholder] = []
                        placeholder_locations[placeholder].append(f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}")
                        if placeholder not in remaining_placeholders:
                            remaining_placeholders.append(placeholder)
        
        # Check headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para_idx, para in enumerate(section.header.paragraphs):
                    text = para.text
                    for match in self.PLACEHOLDER_PATTERN.finditer(text):
                        placeholder = match.group(0)
                        if placeholder not in placeholder_locations:
                            placeholder_locations[placeholder] = []
                        placeholder_locations[placeholder].append(f"header_para_{para_idx}")
                        if placeholder not in remaining_placeholders:
                            remaining_placeholders.append(placeholder)
            
            # Footer
            if section.footer:
                for para_idx, para in enumerate(section.footer.paragraphs):
                    text = para.text
                    for match in self.PLACEHOLDER_PATTERN.finditer(text):
                        placeholder = match.group(0)
                        if placeholder not in placeholder_locations:
                            placeholder_locations[placeholder] = []
                        placeholder_locations[placeholder].append(f"footer_para_{para_idx}")
                        if placeholder not in remaining_placeholders:
                            remaining_placeholders.append(placeholder)
        
        total_count = sum(len(locs) for locs in placeholder_locations.values())
        
        return {
            "valid": len(remaining_placeholders) == 0,
            "remaining_placeholders": remaining_placeholders,
            "placeholder_locations": placeholder_locations,
            "total_placeholders_found": total_count
        }
    
    def _replace_in_content_controls(
        self,
        doc: DocumentType,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> None:
        """
        Replace placeholders in Word content controls (structured document tags).
        
        Content controls are XML elements (w:sdt) that can contain text.
        This method searches for placeholders within content control text content.
        
        Args:
            doc: Document instance
            field_values: Dictionary of placeholder to value mappings
            cdm_data: Optional CDM data for direct path evaluation
        """
        try:
            # Access document XML
            document_xml = doc.part.element
            
            # Find all content controls (w:sdt elements)
            # Content controls use the w:sdt namespace element
            sdt_elements = document_xml.xpath('.//w:sdt', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            for sdt in sdt_elements:
                # Get text content from content control
                # Content control text is typically in w:t elements within w:sdtContent
                text_elements = sdt.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                for text_elem in text_elements:
                    if text_elem.text:
                        original_text = text_elem.text
                        replaced_text = self._replace_placeholders_in_text(
                            original_text, field_values, cdm_data
                        )
                        if replaced_text != original_text:
                            text_elem.text = replaced_text
                            logger.debug(f"Replaced placeholder in content control: {original_text[:50]} -> {replaced_text[:50]}")
        
        except Exception as e:
            # Content controls may not be present in all documents
            # Log but don't fail if they're not found
            logger.debug(f"Could not process content controls (may not be present): {e}")
    
    def _replace_in_textboxes_shapes(
        self,
        doc: DocumentType,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> None:
        """
        Replace placeholders in Word text boxes and shapes.
        
        Text boxes and shapes are stored in drawingML (a:txBody elements).
        This method searches for placeholders within text box and shape text content.
        
        Args:
            doc: Document instance
            field_values: Dictionary of placeholder to value mappings
            cdm_data: Optional CDM data for direct path evaluation
        """
        try:
            # Access document XML
            document_xml = doc.part.element
            
            # Find all text boxes and shapes (a:txBody elements in drawingML)
            # Text boxes use drawingML namespace
            tx_body_elements = document_xml.xpath(
                './/a:txBody',
                namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
            )
            
            for tx_body in tx_body_elements:
                # Get text content from text box/shape
                # Text is in a:p (paragraph) -> a:r (run) -> a:t (text) elements
                text_elements = tx_body.xpath(
                    './/a:t',
                    namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
                )
                
                for text_elem in text_elements:
                    if text_elem.text:
                        original_text = text_elem.text
                        replaced_text = self._replace_placeholders_in_text(
                            original_text, field_values, cdm_data
                        )
                        if replaced_text != original_text:
                            text_elem.text = replaced_text
                            logger.debug(f"Replaced placeholder in text box/shape: {original_text[:50]} -> {replaced_text[:50]}")
        
        except Exception as e:
            # Text boxes/shapes may not be present in all documents
            # Log but don't fail if they're not found
            logger.debug(f"Could not process text boxes/shapes (may not be present): {e}")
    
    def _replace_placeholders_in_text(
        self,
        text: str,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> str:
        """
        Replace placeholders in a text string using the same logic as paragraph replacement.
        
        Args:
            text: Text string containing placeholders
            field_values: Dictionary of placeholder to value mappings
            cdm_data: Optional CDM data for direct path evaluation
            
        Returns:
            Text with placeholders replaced
        """
        result = text
        
        # Replace bracket placeholders [FIELD_NAME]
        for match in self.PLACEHOLDER_PATTERN.finditer(text):
            placeholder = match.group(0)
            field_name = match.group(1)
            
            # Look up in field_values
            if placeholder in field_values:
                result = result.replace(placeholder, str(field_values[placeholder]))
            elif field_name in field_values:
                result = result.replace(placeholder, str(field_values[field_name]))
        
        # Replace double curly placeholders {{field}}
        for match in self.CURLY_PLACEHOLDER_PATTERN.finditer(text):
            placeholder = match.group(0)
            curly_content = match.group(1).strip()
            
            # Normalize to bracket format for lookup
            normalized = self._normalize_curly_placeholder(curly_content)
            bracket_placeholder = f"[{normalized}]"
            
            if bracket_placeholder in field_values:
                result = result.replace(placeholder, str(field_values[bracket_placeholder]))
            elif cdm_data:
                # Try to evaluate as CDM path
                try:
                    cdm_value = self.parser.get_nested_value(cdm_data, curly_content)
                    if cdm_value is not None:
                        formatted_value = self._format_cdm_value(cdm_value)
                        result = result.replace(placeholder, formatted_value)
                except Exception:
                    pass
        
        # Replace single curly placeholders {field}
        # Filter out matches that are inside double curly placeholders
        single_curly_matches = []
        for sc_match in self.SINGLE_CURLY_PLACEHOLDER_PATTERN.finditer(text):
            is_inside_double = False
            for dc_match in self.CURLY_PLACEHOLDER_PATTERN.finditer(text):
                if dc_match.start() <= sc_match.start() and sc_match.end() <= dc_match.end():
                    is_inside_double = True
                    break
            if not is_inside_double:
                single_curly_matches.append(sc_match)
        
        for match in single_curly_matches:
            placeholder = match.group(0)
            curly_content = match.group(1).strip()
            
            if cdm_data:
                try:
                    cdm_value = self.parser.get_nested_value(cdm_data, curly_content)
                    if cdm_value is not None:
                        formatted_value = self._format_cdm_value(cdm_value)
                        result = result.replace(placeholder, formatted_value)
                except Exception:
                    pass
        
        return result
    
    def _format_cdm_value(self, value: Any) -> str:
        """
        Format a CDM value for display in document.
        
        Args:
            value: CDM value (can be various types)
            
        Returns:
            Formatted string representation
        """
        if value is None:
            return ""
        elif isinstance(value, (int, float)):
            return str(value)
        elif isinstance(value, str):
            return value
        elif hasattr(value, 'isoformat'):  # datetime, date
            return value.isoformat()
        elif hasattr(value, '__dict__'):  # Pydantic models
            # For complex objects, return a summary
            if hasattr(value, 'amount') and hasattr(value, 'currency'):
                return f"{value.amount} {value.currency}"
            return str(value)
        else:
            return str(value)
    
    def render_template_to_bytes(
        self,
        template_path: str,
        field_values: Dict[str, str],
        cdm_data: Optional[CreditAgreement] = None
    ) -> bytes:
        """
        Render template and return as bytes.
        
        Args:
            template_path: Path to template Word document
            field_values: Dictionary of placeholder to value mappings
            cdm_data: Optional CDM data for direct path evaluation
            
        Returns:
            Document content as bytes
        """
        doc = self.render_template(template_path, field_values, cdm_data)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
















